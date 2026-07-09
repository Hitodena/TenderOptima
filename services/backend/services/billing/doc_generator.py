"""Generate invoice and act DOCX documents for subscription billing."""

import io
from dataclasses import dataclass
from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from loguru import logger

from backend.core.config import Config
from backend.db.models.subscription_billing import SubscriptionBillingProfile
from backend.services.billing.subscription_lines import BillingQuote
from backend.utils.amount_words import format_byn_amount_line

VAT_RATE = Decimal("20")
VAT_DIVISOR = Decimal("1.20")
UNIT_SERVICE = "услуга"

TABLE_HEADERS = (
    "Наименование услуги",
    "Ед. изм.",
    "Кол-во",
    "Тариф на услугу, без НДС",
    "Стоимость услуги без НДС",
    "Ставка НДС, %",
    "Сумма НДС",
    "Стоимость услуги с НДС",
)

MONTH_NAMES = (
    "",
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)


@dataclass(frozen=True)
class BillingEnvParty:
    """Получатель услуг — реквизиты из BILLING_ISSUER_* (env)."""

    organization_form: str
    organization_name: str
    country: str
    inn: str
    ogrn: str
    legal_address: str
    settlement_account: str
    bank_name: str
    bik: str
    correspondent_account: str


def billing_env_party_from_config(config: Config) -> BillingEnvParty:
    return BillingEnvParty(
        organization_form=config.billing_issuer_organization_form,
        organization_name=config.billing_issuer_name,
        country=config.billing_issuer_country,
        inn=config.billing_issuer_inn,
        ogrn=config.billing_issuer_ogrn,
        legal_address=config.billing_issuer_address,
        settlement_account=config.billing_issuer_settlement_account,
        bank_name=config.billing_issuer_bank_name,
        bik=config.billing_issuer_bik,
        correspondent_account=config.billing_issuer_correspondent_account,
    )


def issuer_from_config(config: Config) -> BillingEnvParty:
    """Alias for billing_env_party_from_config."""
    return billing_env_party_from_config(config)


def _format_doc_date(value: date) -> str:
    month = MONTH_NAMES[value.month]
    return f"«{value.day:02d}» {month} {value.year}г."


def _format_amount(amount: Decimal) -> str:
    return f"{amount:.2f}"


def _split_vat(gross: Decimal) -> tuple[Decimal, Decimal]:
    net = (gross / VAT_DIVISOR).quantize(Decimal("0.01"))
    vat = gross - net
    return net, vat


def _init_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    return doc


def _set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    size: int = 11,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, size=size)


def _organization_title(
    *,
    organization_form: str | None,
    organization_name: str | None,
) -> str:
    if organization_form and organization_name:
        return f'{organization_form} "{organization_name}"'
    if organization_name:
        return organization_name
    if organization_form:
        return organization_form
    return "—"


def _recipient_lines(party: BillingEnvParty) -> list[str]:
    lines = [
        _organization_title(
            organization_form=party.organization_form,
            organization_name=party.organization_name,
        )
    ]
    if party.legal_address:
        lines.append(party.legal_address)
    if party.settlement_account:
        lines.append(f"Р/с {party.settlement_account}")
    if party.bank_name:
        lines.append(party.bank_name)
    if party.bik:
        lines.append(f"БИК {party.bik}")
    if party.inn:
        lines.append(f"УНП {party.inn}")
    if party.ogrn:
        lines.append(f"ОКПО {party.ogrn}")
    return lines


def _payer_lines(profile: SubscriptionBillingProfile) -> list[str]:
    lines = [
        _organization_title(
            organization_form=profile.organization_form,
            organization_name=profile.organization_name,
        )
    ]
    if profile.legal_address:
        lines.append(profile.legal_address)
    if profile.inn:
        lines.append(f"УНП {profile.inn}")
    if profile.ogrn:
        lines.append(f"ОКПО {profile.ogrn}")
    return lines


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    size: int = 9,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = 1


def _add_vat_items_table(
    doc: Document,
    quote: BillingQuote,
) -> tuple[Decimal, Decimal, Decimal]:
    currency_paragraph = doc.add_paragraph()
    currency_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    currency_run = currency_paragraph.add_run(quote.currency_code)
    _set_run_font(currency_run, bold=True)

    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0].cells
    for index, title in enumerate(TABLE_HEADERS):
        _set_cell_text(header_row[index], title, bold=True, size=8)

    total_net = Decimal("0")
    total_vat = Decimal("0")
    total_gross = Decimal("0")

    for item in quote.line_items:
        quantity = Decimal("1")
        gross = item.amount
        net, vat = _split_vat(gross)
        tariff_net = net / quantity

        total_net += net
        total_vat += vat
        total_gross += gross

        row = table.add_row().cells
        _set_cell_text(
            row[0], item.name, align=WD_ALIGN_PARAGRAPH.LEFT, size=8
        )
        _set_cell_text(row[1], UNIT_SERVICE, size=8)
        _set_cell_text(row[2], str(int(quantity)), size=8)
        _set_cell_text(row[3], _format_amount(tariff_net), size=8)
        _set_cell_text(row[4], _format_amount(net), size=8)
        _set_cell_text(row[5], str(int(VAT_RATE)), size=8)
        _set_cell_text(row[6], _format_amount(vat), size=8)
        _set_cell_text(row[7], _format_amount(gross), size=8)

    totals_row = table.add_row().cells
    _set_cell_text(totals_row[0], "Итого", bold=True, size=8)
    _set_cell_text(totals_row[1], "", size=8)
    _set_cell_text(totals_row[2], "", size=8)
    _set_cell_text(totals_row[3], "", size=8)
    _set_cell_text(totals_row[4], _format_amount(total_net), bold=True, size=8)
    _set_cell_text(totals_row[5], "", size=8)
    _set_cell_text(totals_row[6], _format_amount(total_vat), bold=True, size=8)
    _set_cell_text(
        totals_row[7],
        _format_amount(total_gross),
        bold=True,
        size=8,
    )

    return total_net, total_vat, total_gross


def _add_amount_summary(
    doc: Document, total_gross: Decimal, total_vat: Decimal
) -> None:
    doc.add_paragraph()
    _add_paragraph(doc, f"Итого: {format_byn_amount_line(total_gross)}")
    _add_paragraph(
        doc,
        (
            f"В том числе НДС {int(VAT_RATE)}%: "
            f"{format_byn_amount_line(total_vat)}"
        ),
    )


def _add_invoice_header(
    doc: Document,
    *,
    receipt_id: str,
    doc_date: date,
    recipient: BillingEnvParty,
    payer: SubscriptionBillingProfile,
) -> None:
    header = doc.add_table(rows=1, cols=2)
    header.autofit = True
    left, right = header.rows[0].cells

    left_paragraph = left.paragraphs[0]
    title_run = left_paragraph.add_run("Получатель\n")
    _set_run_font(title_run, bold=True)
    for line in _recipient_lines(recipient):
        line_run = left.add_paragraph().add_run(line)
        _set_run_font(line_run)

    payer_title = left.add_paragraph()
    payer_title_run = payer_title.add_run("Плательщик")
    _set_run_font(payer_title_run, bold=True)
    for line in _payer_lines(payer):
        line_run = left.add_paragraph().add_run(line)
        _set_run_font(line_run)

    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    invoice_run = right_paragraph.add_run(f"СЧЕТ-ФАКТУРА № {receipt_id}\n")
    _set_run_font(invoice_run, bold=True, size=12)
    date_run = right_paragraph.add_run(f"от {_format_doc_date(doc_date)}")
    _set_run_font(date_run, size=11)

    doc.add_paragraph()


def build_invoice_docx_bytes(
    *,
    quote: BillingQuote,
    service_recipient: SubscriptionBillingProfile,
    env_party: BillingEnvParty,
) -> bytes:
    doc = _init_document()
    doc_date = date.today()

    _add_invoice_header(
        doc,
        receipt_id=quote.receipt_id,
        doc_date=doc_date,
        recipient=env_party,
        payer=service_recipient,
    )
    _, total_vat, total_gross = _add_vat_items_table(doc, quote)
    _add_amount_summary(doc, total_gross, total_vat)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_act_docx_bytes(
    *,
    quote: BillingQuote,
    service_recipient: SubscriptionBillingProfile,
    env_party: BillingEnvParty,
) -> bytes:
    del env_party  # act body does not include party footer blocks
    doc = _init_document()
    doc_date = date.today()
    period = (
        f"{quote.period_start.strftime('%d.%m.%Y')} — "
        f"{quote.period_end.date().strftime('%d.%m.%Y')}"
    )

    _add_paragraph(
        doc,
        f"АКТ СДАЧИ-ПРИЕМКИ УСЛУГ № {quote.receipt_id}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
    )
    _add_paragraph(
        doc,
        "публичному договору возмездного оказания услуг",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    location_date = doc.add_table(rows=1, cols=2)
    left_cell, right_cell = location_date.rows[0].cells
    left_run = left_cell.paragraphs[0].add_run("г. Минск")
    _set_run_font(left_run)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = right_paragraph.add_run(_format_doc_date(doc_date))
    _set_run_font(date_run)

    doc.add_paragraph()
    payer_name = _organization_title(
        organization_form=service_recipient.organization_form,
        organization_name=service_recipient.organization_name,
    )
    _add_paragraph(
        doc,
        (
            "Настоящий акт составлен о том, что услуги по подписке "
            f"«{quote.plan_title}» за период {period} оказаны "
            f"в полном объёме для {payer_name}."
        ),
    )
    doc.add_paragraph()

    _, total_vat, total_gross = _add_vat_items_table(doc, quote)
    _add_amount_summary(doc, total_gross, total_vat)
    doc.add_paragraph()
    _add_paragraph(
        doc,
        (
            "Заказчик претензий по объёму, качеству и срокам оказания "
            "услуг не имеет. Настоящий акт составлен в двух экземплярах, "
            "имеющих одинаковую юридическую силу."
        ),
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def write_billing_documents(
    *,
    quote: BillingQuote,
    service_recipient: SubscriptionBillingProfile,
    env_party: BillingEnvParty,
    invoice_path: Path,
    act_path: Path,
) -> tuple[Path, Path]:
    """Build invoice/act as DOCX, convert to PDF, return PDF paths."""
    from backend.services.extraction.docx_to_pdf import (
        convert_docx_to_pdf_file,
    )

    invoice_path.parent.mkdir(parents=True, exist_ok=True)
    invoice_docx = invoice_path.with_suffix(".docx")
    act_docx = act_path.with_suffix(".docx")
    invoice_pdf = invoice_path.with_suffix(".pdf")
    act_pdf = act_path.with_suffix(".pdf")

    invoice_docx.write_bytes(
        build_invoice_docx_bytes(
            quote=quote,
            service_recipient=service_recipient,
            env_party=env_party,
        )
    )
    act_docx.write_bytes(
        build_act_docx_bytes(
            quote=quote,
            service_recipient=service_recipient,
            env_party=env_party,
        )
    )

    converted_invoice = convert_docx_to_pdf_file(invoice_docx, invoice_pdf)
    converted_act = convert_docx_to_pdf_file(act_docx, act_pdf)
    if converted_invoice is None or converted_act is None:
        raise RuntimeError(
            "Failed to convert billing DOCX to PDF "
            "(LibreOffice/soffice required)"
        )

    logger.info(
        "Billing documents written",
        invoice_docx=str(invoice_docx),
        act_docx=str(act_docx),
        invoice_pdf=str(converted_invoice),
        act_pdf=str(converted_act),
    )
    return converted_invoice, converted_act
