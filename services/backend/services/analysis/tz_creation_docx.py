"""DOCX export for the TZ creation wizard (Module 3).

Unlike ``docx_export.py`` (supplier clarification letters), this module
renders a technical specification outline as a numbered hierarchical
table: ``1 -> 1.1, 1.2 -> 1.2.1, 1.2.2`` etc, using the same
``RequirementNode`` hierarchy shape as Module 2's ``requirements_tz``.
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.styles.style import ParagraphStyle

from backend.services.analysis.docx_export import format_ru_date
from backend.utils.requirements_struct import (
    RequirementNode,
    iter_hierarchy_rows,
)

DEFAULT_TITLE = "Тендер на закупку"
_NUMBER_COLUMN_WIDTH = Cm(2.5)
_TEXT_COLUMN_WIDTH = Cm(14.5)
_INDENT_PER_LEVEL = Cm(0.5)


def _set_normal_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    if isinstance(normal, ParagraphStyle):
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)


def _add_title(doc: Document, title: str) -> None:
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title.strip() or DEFAULT_TITLE)
    run.font.name = "Times New Roman"


def _add_header_row(table) -> None:
    header_cells = table.rows[0].cells
    header_cells[0].text = "№ п/п"
    header_cells[1].text = "Наименование раздела / требование"
    for cell in header_cells:
        cell.width = (
            _NUMBER_COLUMN_WIDTH
            if cell is header_cells[0]
            else _TEXT_COLUMN_WIDTH
        )
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True


def _add_requirement_row(
    table,
    key: str,
    depth: int,
    node: RequirementNode,
) -> None:
    row_cells = table.add_row().cells
    row_cells[0].width = _NUMBER_COLUMN_WIDTH
    row_cells[1].width = _TEXT_COLUMN_WIDTH

    number_paragraph = row_cells[0].paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = number_paragraph.add_run(key)
    number_run.bold = depth == 0

    text_paragraph = row_cells[1].paragraphs[0]
    text_paragraph.paragraph_format.left_indent = _INDENT_PER_LEVEL * depth
    text_run = text_paragraph.add_run(node["text"])
    text_run.bold = depth == 0


def build_tz_creation_docx(
    title: str,
    hierarchy: dict[str, RequirementNode],
) -> bytes:
    """Render a TZ outline as a numbered hierarchical table in DOCX.

    ``title`` defaults to "Тендер на закупку" when blank. Rows follow the
    same document order as the requirement keys (``1 -> 1.1 -> 1.2 ->
    1.2.1 -> 1.2.2 -> 2 -> ...``).
    """
    doc = Document()
    _set_normal_style(doc)
    _add_title(doc, title)

    rows = iter_hierarchy_rows(hierarchy)

    if not rows:
        doc.add_paragraph("Требования пока не добавлены.")
    else:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        _add_header_row(table)
        for key, depth, node in rows:
            _add_requirement_row(table, key, depth, node)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Дата формирования: {format_ru_date()}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
