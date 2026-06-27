import io
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.styles.style import ParagraphStyle

from backend.enums import TZAnalysisStatus
from backend.schemas.analysis import TZAnalysisItem
from backend.utils.requirements_struct import (
    format_tz_requirement_ref,
    requirement_key_from_flat,
)

_RU_MONTHS_GENITIVE = (
    "",
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)

HEADER_MISMATCH = "Не соответствует:"
HEADER_NOT_FOUND = "Не найдено:"
HEADER_PARTIAL = "Требуют уточнения/дополнения:"


def format_ru_date(value: date | None = None) -> str:
    """Format date as '17 июня 2026 г.' for letter footers."""
    current = value or date.today()
    month = _RU_MONTHS_GENITIVE[current.month]
    return f"{current.day} {month} {current.year} г."


def _mismatch_reason(item: TZAnalysisItem) -> str:
    if item.status == TZAnalysisStatus.NOT_FOUND:
        return f"Параметр не найден: {item.explanation}"
    return f"Причина отклонения: {item.explanation}"


def _letter_tz_key(item: TZAnalysisItem) -> str | None:
    if item.ref:
        return item.ref
    return requirement_key_from_flat(item.requirement)


def _letter_tz_requirement_text(item: TZAnalysisItem) -> str:
    if item.ref_value:
        key = _letter_tz_key(item)
        if key:
            return f"{key}. {item.ref_value}"
        return item.ref_value
    return item.requirement


def _letter_tz_requirement_ref(item: TZAnalysisItem) -> str | None:
    key = _letter_tz_key(item)
    text = item.ref_value
    if key and text:
        return format_tz_requirement_ref(
            key, None, f"{key}. {text}", quote=text
        )
    return item.requirement_ref


def build_clarification_docx_from_paragraphs(paragraphs: list[str]) -> bytes:
    """Build DOCX from user-edited letter paragraphs."""
    doc = Document()
    normal = doc.styles["Normal"]
    if isinstance(normal, ParagraphStyle):
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

    title = "О несоответствии предложения техническим требованиям"
    bold_headers = {
        HEADER_MISMATCH,
        HEADER_NOT_FOUND,
        HEADER_PARTIAL,
    }

    for index, text in enumerate(paragraphs):
        if text == "":
            doc.add_paragraph()
            continue
        if index == 0 and text == title:
            paragraph = doc.add_paragraph()
            paragraph.add_run(text).bold = True
            continue
        if text in bold_headers:
            paragraph = doc.add_paragraph()
            paragraph.add_run(text).bold = True
            continue
        doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _append_mismatch_items(
    doc: Document,
    items: list[TZAnalysisItem],
    start_num: int,
) -> int:
    num = start_num
    for item in items:
        doc.add_paragraph(f"{num}. По пункту:")
        doc.add_paragraph(
            f'Требование: "{_letter_tz_requirement_text(item)}"'
            + (
                f" ({ref})"
                if (ref := _letter_tz_requirement_ref(item))
                else ""
            )
        )
        if item.offer_value:
            doc.add_paragraph(
                f'Предложено: "{item.offer_value}"'
                + (f" ({item.offer_ref})" if item.offer_ref else "")
            )
        doc.add_paragraph(_mismatch_reason(item))
        num += 1
    return num


def _append_not_found_items(
    doc: Document,
    items: list[TZAnalysisItem],
    start_num: int,
) -> int:
    num = start_num
    for item in items:
        doc.add_paragraph(f"{num}. По пункту:")
        doc.add_paragraph(
            f'Требование: "{_letter_tz_requirement_text(item)}"'
            + (
                f" ({ref})"
                if (ref := _letter_tz_requirement_ref(item))
                else ""
            )
        )
        doc.add_paragraph(_mismatch_reason(item))
        num += 1
    return num


def build_clarification_docx(
    items: list[TZAnalysisItem],
    selected_indices: list[int],
    deadline_date: str | None = None,
) -> bytes:
    """Build a DOCX letter requesting clarifications from the supplier."""
    doc = Document()
    normal = doc.styles["Normal"]
    if isinstance(normal, ParagraphStyle):
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

    title = doc.add_paragraph()
    title.add_run(
        "О несоответствии предложения техническим требованиям"
    ).bold = True

    doc.add_paragraph(
        "Проведён анализ вашего предложения по соответствию "
        "техническому заданию."
    )
    doc.add_paragraph("Выявлены следующие замечания и требуемые уточнения:")

    selected = [
        items[i]
        for i in selected_indices
        if 0 <= i < len(items)
        and items[i].status
        in (
            TZAnalysisStatus.PARTIAL,
            TZAnalysisStatus.MISSING,
            TZAnalysisStatus.NOT_FOUND,
        )
    ]

    mismatches = [
        it for it in selected if it.status == TZAnalysisStatus.MISSING
    ]
    not_found = [
        it for it in selected if it.status == TZAnalysisStatus.NOT_FOUND
    ]
    clarifications = [
        it for it in selected if it.status == TZAnalysisStatus.PARTIAL
    ]

    item_num = 1

    if mismatches:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(HEADER_MISMATCH).bold = True
        item_num = _append_mismatch_items(doc, mismatches, item_num)

    if not_found:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(HEADER_NOT_FOUND).bold = True
        item_num = _append_not_found_items(doc, not_found, item_num)

    if clarifications:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(HEADER_PARTIAL).bold = True
        for item in clarifications:
            doc.add_paragraph(f"{item_num}. Пункт:")
            doc.add_paragraph(
                f'Требуется: "{_letter_tz_requirement_text(item)}"'
                + (
                    f" ({ref})"
                    if (ref := _letter_tz_requirement_ref(item))
                    else ""
                )
            )
            if item.offer_value:
                doc.add_paragraph(
                    f'Предложено: "{item.offer_value}"'
                    + (f" ({item.offer_ref})" if item.offer_ref else "")
                )
            doc.add_paragraph(f"Необходимо: {item.explanation}")
            item_num += 1

    deadline = deadline_date or "7 дней"
    doc.add_paragraph()
    doc.add_paragraph(
        f"Просим предоставить дополненное/уточненное предложение "
        f"не позднее {deadline}."
    )
    doc.add_paragraph()
    doc.add_paragraph(format_ru_date())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _append_preview_mismatch_lines(
    paragraphs: list[str],
    items: list[TZAnalysisItem],
    start_num: int,
) -> int:
    num = start_num
    for item in items:
        paragraphs.append(f"{num}. По пункту:")
        req = f'Требование: "{_letter_tz_requirement_text(item)}"'
        if ref := _letter_tz_requirement_ref(item):
            req += f" ({ref})"
        paragraphs.append(req)
        if item.offer_value:
            offer = f'Предложено: "{item.offer_value}"'
            if item.offer_ref:
                offer += f" ({item.offer_ref})"
            paragraphs.append(offer)
        paragraphs.append(_mismatch_reason(item))
        num += 1
    return num


def _append_preview_not_found_lines(
    paragraphs: list[str],
    items: list[TZAnalysisItem],
    start_num: int,
) -> int:
    num = start_num
    for item in items:
        paragraphs.append(f"{num}. По пункту:")
        req = f'Требование: "{_letter_tz_requirement_text(item)}"'
        if ref := _letter_tz_requirement_ref(item):
            req += f" ({ref})"
        paragraphs.append(req)
        paragraphs.append(_mismatch_reason(item))
        num += 1
    return num


def build_clarification_preview(
    items: list[TZAnalysisItem],
    selected_indices: list[int],
    deadline_date: str | None = None,
) -> tuple[str, list[str], bool]:
    """Build plain-text preview paragraphs for the clarification letter."""
    selected = [
        items[i]
        for i in selected_indices
        if 0 <= i < len(items)
        and items[i].status
        in (
            TZAnalysisStatus.PARTIAL,
            TZAnalysisStatus.MISSING,
            TZAnalysisStatus.NOT_FOUND,
        )
    ]

    mismatches = [
        it for it in selected if it.status == TZAnalysisStatus.MISSING
    ]
    not_found = [
        it for it in selected if it.status == TZAnalysisStatus.NOT_FOUND
    ]
    clarifications = [
        it for it in selected if it.status == TZAnalysisStatus.PARTIAL
    ]
    has_issues = bool(mismatches or not_found or clarifications)

    paragraphs: list[str] = [
        "О несоответствии предложения техническим требованиям",
        "",
        "Проведён анализ вашего предложения по соответствию "
        "техническому заданию.",
        "Выявлены следующие замечания и требуемые уточнения:",
    ]

    item_num = 1

    if mismatches:
        paragraphs.extend(["", HEADER_MISMATCH])
        item_num = _append_preview_mismatch_lines(
            paragraphs, mismatches, item_num
        )

    if not_found:
        paragraphs.extend(["", HEADER_NOT_FOUND])
        item_num = _append_preview_not_found_lines(
            paragraphs, not_found, item_num
        )

    if clarifications:
        paragraphs.extend(["", HEADER_PARTIAL])
        for item in clarifications:
            paragraphs.append(f"{item_num}. Пункт:")
            req = f'Требуется: "{_letter_tz_requirement_text(item)}"'
            if ref := _letter_tz_requirement_ref(item):
                req += f" ({ref})"
            paragraphs.append(req)
            if item.offer_value:
                offer = f'Предложено: "{item.offer_value}"'
                if item.offer_ref:
                    offer += f" ({item.offer_ref})"
                paragraphs.append(offer)
            paragraphs.append(f"Необходимо: {item.explanation}")
            item_num += 1

    deadline = deadline_date or "7 дней"
    paragraphs.extend(
        [
            "",
            f"Просим предоставить дополненное/уточненное предложение "
            f"не позднее {deadline}.",
            "",
            format_ru_date(),
        ]
    )

    return (
        "О несоответствии предложения техническим требованиям",
        paragraphs,
        has_issues,
    )
