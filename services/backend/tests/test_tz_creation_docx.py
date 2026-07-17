"""Unit tests for iter_hierarchy_rows and the TZ creation .docx exporter."""

import io

from backend.services.analysis.tz_creation_docx import (
    DEFAULT_TITLE,
    build_tz_creation_docx,
)
from backend.utils.requirements_struct import iter_hierarchy_rows
from docx import Document

HIERARCHY = {
    "1": {
        "text": "Общие требования",
        "children": {
            "1.1": {"text": "Пункт 1.1", "children": {}},
            "1.2": {
                "text": "Пункт 1.2",
                "children": {
                    "1.2.1": {"text": "Подпункт 1.2.1", "children": {}},
                    "1.2.2": {"text": "Подпункт 1.2.2", "children": {}},
                },
            },
        },
    },
    "2": {"text": "Технические характеристики", "children": {}},
}


def test_iter_hierarchy_rows_follows_document_order():
    rows = iter_hierarchy_rows(HIERARCHY)
    assert [(key, depth) for key, depth, _ in rows] == [
        ("1", 0),
        ("1.1", 1),
        ("1.2", 1),
        ("1.2.1", 2),
        ("1.2.2", 2),
        ("2", 0),
    ]


def test_iter_hierarchy_rows_empty_input():
    assert iter_hierarchy_rows({}) == []


def _read_table_rows(docx_bytes: bytes) -> list[tuple[str, str]]:
    doc = Document(io.BytesIO(docx_bytes))
    table = doc.tables[0]
    return [
        (row.cells[0].text.strip(), row.cells[1].text.strip())
        for row in table.rows
    ]


def test_build_tz_creation_docx_default_title_when_blank():
    docx_bytes = build_tz_creation_docx("", HIERARCHY)
    doc = Document(io.BytesIO(docx_bytes))
    assert doc.paragraphs[0].text == DEFAULT_TITLE


def test_build_tz_creation_docx_uses_custom_title():
    docx_bytes = build_tz_creation_docx(
        "Тендер на закупку оборудования", HIERARCHY
    )
    doc = Document(io.BytesIO(docx_bytes))
    assert doc.paragraphs[0].text == "Тендер на закупку оборудования"


def test_build_tz_creation_docx_table_matches_hierarchy_order():
    docx_bytes = build_tz_creation_docx(DEFAULT_TITLE, HIERARCHY)
    rows = _read_table_rows(docx_bytes)

    assert rows[0] == ("№ п/п", "Наименование раздела / требование")
    assert rows[1] == ("1", "Общие требования")
    assert rows[2] == ("1.1", "Пункт 1.1")
    assert rows[3] == ("1.2", "Пункт 1.2")
    assert rows[4] == ("1.2.1", "Подпункт 1.2.1")
    assert rows[5] == ("1.2.2", "Подпункт 1.2.2")
    assert rows[6] == ("2", "Технические характеристики")


def test_build_tz_creation_docx_empty_hierarchy_still_produces_a_document():
    docx_bytes = build_tz_creation_docx("Пустое ТЗ", {})
    doc = Document(io.BytesIO(docx_bytes))
    assert not doc.tables
    assert any(
        "Требования пока не добавлены" in p.text for p in doc.paragraphs
    )
