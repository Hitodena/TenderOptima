#!/usr/bin/env python3
"""
OPTIMIZED Section-Aware Text Extraction for SupplierFinder Analysis
Enhanced text extraction with improved section detection, proper numerical sorting, and speed optimizations
"""

import sys
import os
import re
import tempfile
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    from docx import Document
    import pandas as pd
    import openpyxl
except ImportError as e:
    print(f"Missing required package: {e}", file=sys.stderr)
    sys.exit(1)

def detect_sections_optimized(text):
    """OPTIMIZED: Detect document sections with enhanced patterns and proper sorting"""
    sections = {}
    current_section = "general"
    
    # ENHANCED section patterns - MORE COMPREHENSIVE DETECTION
    section_patterns = [
        # PRIMARY: Multi-level numbered sections (4.1.2, 1.3.5, etc.)
        (r'^\s*(\d+)\.(\d+)\.(\d+)\s+([А-ЯЁA-Z][а-яё\s,\-a-zA-Z]{3,})', 'numbered_subsection_L3'),
        (r'^\s*(\d+)\.(\d+)\s+([А-ЯЁA-Z][а-яё\s,\-a-zA-Z]{3,})', 'numbered_subsection_L2'),
        (r'^\s*(\d+)\.\s*([А-ЯЁA-Z][а-яё\s,\-a-zA-Z]{3,})', 'numbered_section_L1'),
        
        # SECONDARY: Section number only (for capturing orphaned numbers)
        (r'^\s*(\d+)\.(\d+)\.(\d+)\s*$', 'section_number_L3'),
        (r'^\s*(\d+)\.(\d+)\s*$', 'section_number_L2'),
        (r'^\s*(\d+)\.\s*$', 'section_number_L1'),
        
        # TERTIARY: Formal section titles
        (r'^\s*(Раздел\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'section_title'),
        (r'^\s*(Глава\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'chapter_title'),
        
        # QUATERNARY: Technical content patterns
        (r'^\s*(Техническ[а-я]*\s+характеристик[а-я]*)', 'technical_specifications'),
        (r'^\s*(Комплект[а-я]*\s+оборудован[а-я]*)', 'equipment'),
        (r'^\s*(Комплект[а-я]*\s+технологическ[а-я]*)', 'equipment'),
        (r'^\s*(Систем[а-я]*\s+автоматизац[а-я]*)', 'automation_control'),
        (r'^\s*(Систем[а-я]*\s+контрол[а-я]*)', 'automation_control'),
        (r'^\s*(Общ[а-я]*\s+требован[а-я]*)', 'general_requirements'),
        (r'^\s*(Безопасност[а-я]*)', 'safety'),
        
        # ENGLISH patterns
        (r'^\s*(\d+)\.(\d+)\.(\d+)\s+([A-Z][a-z\s,\-]{3,})', 'numbered_subsection_L3_en'),
        (r'^\s*(\d+)\.(\d+)\s+([A-Z][a-z\s,\-]{3,})', 'numbered_subsection_L2_en'),
        (r'^\s*(\d+)\.\s*([A-Z][a-z\s,\-]{3,})', 'numbered_section_L1_en'),
        (r'^\s*(Technical\s+Specifications?)', 'technical_specifications'),
        (r'^\s*(Equipment\s+Set)', 'equipment'),
        (r'^\s*(Automation\s+System)', 'automation_control'),
    ]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Check for section headers - OPTIMIZED PATTERN MATCHING
        section_detected = False
        for pattern, section_type in section_patterns:
            match = re.match(pattern, line_stripped, re.IGNORECASE | re.MULTILINE)
            if match:
                section_title = match.group(0).strip()
                
                # ENHANCED section categorization with hierarchical numbering
                if 'numbered' in section_type:
                    # Extract all numeric groups for hierarchical sections
                    groups = [g for g in match.groups() if g and g.isdigit()]
                    if groups:
                        # Create hierarchical section identifier (e.g., "4_1_2")
                        current_section = f"section_{'_'.join(groups)}"
                    else:
                        current_section = f"section_{len(sections) + 1}"
                        
                elif 'section_number' in section_type:
                    # Handle orphaned section numbers
                    groups = [g for g in match.groups() if g and g.isdigit()]
                    if groups:
                        current_section = f"section_{'_'.join(groups)}"
                    else:
                        current_section = f"section_{len(sections) + 1}"
                        
                elif any(word in line_stripped.lower() for word in ['техническ', 'характеристик', 'technical', 'specification']):
                    current_section = "technical_specifications"
                elif any(word in line_stripped.lower() for word in ['комплект', 'оборудован', 'equipment', 'технологическ']):
                    current_section = "equipment"
                elif any(word in line_stripped.lower() for word in ['автоматизац', 'automation', 'контрол', 'control']):
                    current_section = "automation_control"
                elif any(word in line_stripped.lower() for word in ['безопасност', 'safety']):
                    current_section = "safety"
                elif any(word in line_stripped.lower() for word in ['общ', 'требован', 'general', 'requirement']):
                    current_section = "general_requirements"
                else:
                    current_section = f"section_{len(sections) + 1}"
                
                if current_section not in sections:
                    sections[current_section] = {
                        'title': section_title,
                        'content': [],
                        'start_line': i,
                        'section_type': section_type,
                        'numeric_order': extract_numeric_order(current_section)
                    }
                
                section_detected = True
                break
        
        # Add content to current section
        if current_section not in sections:
            sections[current_section] = {
                'title': 'General Content',
                'content': [],
                'start_line': 0,
                'section_type': 'general',
                'numeric_order': [0]
            }
        
        sections[current_section]['content'].append(line)
    
    return sections

def extract_numeric_order(section_name):
    """Extract numeric ordering from section name for proper sorting"""
    if section_name.startswith('section_'):
        numbers = section_name.replace('section_', '').split('_')
        try:
            return [int(n) for n in numbers if n.isdigit()]
        except ValueError:
            return [999]  # Put non-numeric at end
    return [0]  # General sections first

def sort_sections_properly(sections):
    """Sort sections by proper numerical order (1, 2, 3, 4, 5, 6, 7...)"""
    def sort_key(item):
        section_name, section_data = item
        return section_data.get('numeric_order', [999])
    
    return sorted(sections.items(), key=sort_key)

def extract_text_from_pdf_optimized(file_path):
    """OPTIMIZED: Extract text from PDF with enhanced section detection and speed improvements"""
    text = ""
    
    try:
        # SPEED OPTIMIZATION: Use faster text extraction first
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Process pages in batches for memory efficiency
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n[PAGE:{page_num}]\n{page_text}\n"
                except Exception as e:
                    print(f"Error extracting page {page_num}: {e}", file=sys.stderr)
                    continue
        
        # OCR fallback only if no text extracted
        if not text.strip():
            print("No text found via PyPDF2, falling back to OCR...", file=sys.stderr)
            images = convert_from_path(file_path)
            for page_num, image in enumerate(images, 1):
                ocr_text = pytesseract.image_to_string(image, lang='rus+eng')
                if ocr_text.strip():
                    text += f"\n[PAGE:{page_num}]\n{ocr_text}\n"
    
    except Exception as e:
        print(f"Error extracting from PDF: {e}", file=sys.stderr)
        return ""
    
    # ENHANCED section detection with proper sorting
    if text.strip():
        sections = detect_sections_optimized(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_docx_optimized(file_path):
    """OPTIMIZED: Extract text from DOCX with enhanced section detection"""
    text = ""
    page_num = 1
    
    try:
        doc = Document(file_path)
        
        # SPEED OPTIMIZATION: Batch process paragraphs
        paragraph_texts = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                paragraph_texts.append(para_text)
        
        # Add page breaks every 50 paragraphs (estimation)
        for i, para_text in enumerate(paragraph_texts):
            if i > 0 and i % 50 == 0:
                page_num += 1
                text += f"\n[PAGE:{page_num}]\n"
            text += para_text + "\n"
        
        # Process tables efficiently
        for table in doc.tables:
            text += f"\n[TABLE]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += f"[/TABLE]\n"
        
    except Exception as e:
        print(f"Error extracting from DOCX: {e}", file=sys.stderr)
        return ""
    
    # ENHANCED section detection with proper sorting
    if text.strip():
        sections = detect_sections_optimized(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_txt_optimized(file_path):
    """OPTIMIZED: Extract text from plain text files with enhanced section detection"""
    try:
        # SPEED OPTIMIZATION: Try UTF-8 first (most common)
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='cp1251') as file:
                text = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='iso-8859-1') as file:
                text = file.read()
    
    # ENHANCED section detection with proper sorting
    if text.strip():
        sections = detect_sections_optimized(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_text_optimized.py <file_path> <mime_type>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    mime_type = sys.argv[2]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    try:
        if mime_type.startswith('application/pdf'):
            text = extract_text_from_pdf_optimized(file_path)
        elif mime_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            text = extract_text_from_docx_optimized(file_path)
        elif mime_type.startswith('text/'):
            text = extract_text_from_txt_optimized(file_path)
        else:
            # Fallback based on file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.pdf':
                text = extract_text_from_pdf_optimized(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = extract_text_from_docx_optimized(file_path)
            elif file_ext == '.txt':
                text = extract_text_from_txt_optimized(file_path)
            else:
                print(f"Unsupported file type: {mime_type}", file=sys.stderr)
                sys.exit(1)
        
        print(text)
    
    except Exception as e:
        print(f"Error processing file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()