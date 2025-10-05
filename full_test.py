# ==============================================================================
#           КОМПЛЕКСНЫЙ ПЛАН ТЕСТИРОВАНИЯ И ВНЕДРЕНИЯ
# ==============================================================================
#
# Скопируйте этот код в один файл (например, full_test.py) 
# и запустите его командой: python full_test.py
#
# Не забудьте установить все зависимости:
# pip install PyPDF2 python-docx pandas openpyxl easyocr pdf2image torch torchvision torchaudio docx2txt Pillow
#
# Также установите 'antiword' и 'poppler' в вашей системе.
#

import os
import subprocess
import pandas as pd
import docx
from PIL import Image, ImageDraw, ImageFont
import PyPDF2
from pdf2image import convert_from_path
import easyocr
import docx2txt

# --- ЗАДАЧА 1: ИНИЦИАЛИЗАЦИЯ OCR ---
# ------------------------------------
# Инициализируем EasyOCR ридер один раз для экономии ресурсов
try:
    reader = easyocr.Reader(['ru', 'en']) 
    print("INFO: EasyOCR ридер успешно инициализирован.")
except Exception as e:
    print(f"FATAL: Не удалось инициализировать EasyOCR: {e}")
    reader = None

# --- ЗАДАЧА 2: ВСЕ ФУНКЦИИ ДЛЯ ИЗВЛЕЧЕНИЯ ТЕКСТА ---
# ----------------------------------------------------

def extract_text_from_txt(path):
    # Эта функция для простых текстовых файлов
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"ERROR: {e}"

def extract_text_from_docx(path):
    # Используем docx2txt, так как он проще и надежнее для простого извлечения текста
    try:
        return docx2txt.process(path)
    except Exception as e:
        return f"ERROR: {e}"

def extract_text_from_excel(path):
    # Используем pandas для .xlsx и .xls
    try:
        df = pd.read_excel(path, sheet_name=None) # Читаем все листы
        full_text = ""
        for sheet_name, sheet_df in df.items():
            full_text += f"--- Sheet: {sheet_name} ---\n{sheet_df.to_string()}\n\n"
        return full_text
    except Exception as e:
        return f"ERROR: {e}"

def extract_text_from_csv(path):
    # Исправленная функция для CSV
    try:
        df = pd.read_csv(path, encoding='utf-8')
        return df.to_string()
    except UnicodeDecodeError:
        try:
            df = pd.read_csv(path, encoding='cp1251')
            return df.to_string()
        except Exception as e:
            return f"ERROR: {e}"
    except Exception as e:
        return f"ERROR: {e}"

def extract_text_from_image(path):
    # OCR для отдельных изображений
    if not reader:
        return "ERROR: EasyOCR не инициализирован."
    try:
        result = reader.readtext(path, detail=0, paragraph=True)
        return " ".join(result)
    except Exception as e:
        return f"ERROR: {e}"

def extract_text_from_scanned_pdf(path):
    # Функция для PDF, которая сначала проверяет текст, потом делает OCR
    if not reader:
        return "ERROR: EasyOCR не инициализирован."
    try:
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
            if len(text.strip()) > 100:
                return text.strip()
    except Exception:
        pass # Игнорируем ошибки PyPDF2 и переходим к OCR
    
    # Если текста нет, делаем OCR
    try:
        images = convert_from_path(path)
        full_text_ocr = ""
        for i, image in enumerate(images):
            result = reader.readtext(image, detail=0, paragraph=True)
            full_text_ocr += " ".join(result) + f"\n\n--- Page {i+1} ---\n\n"
        return full_text_ocr.strip()
    except Exception as e:
        return f"ERROR: OCR для PDF не удался. Ошибка: {e}"

def extract_text_from_doc(doc_path):
    """
    Извлекает текст из .doc файлов.
    Основной метод: 'antiword' для максимальной совместимости.
    Резервный метод (Fallback): 'docx2txt' на случай, если antiword не установлен.
    """
    if not os.path.exists(doc_path):
        return f"ERROR: Файл не найден: {doc_path}"

    # --- Метод 1: Попытка использовать 'antiword' (предпочтительный) ---
    try:
        # Проверяем наличие antiword без вывода ошибок в консоль
        subprocess.run(['antiword', '-h'], check=True, capture_output=True)
        
        # Если antiword найден, используем его
        process = subprocess.run(
            ['antiword', doc_path],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore'
        )
        if process.returncode == 0 and process.stdout.strip():
            print(f"INFO: Текст из {doc_path} успешно извлечен с помощью antiword.")
            return process.stdout.strip()
        else:
            print(f"WARNING: antiword не смог обработать {doc_path}. Переключаюсь на fallback.")

    except (FileNotFoundError, subprocess.CalledProcessError):
        # Эта ошибка означает, что antiword не найден или не работает
        print(f"INFO: Утилита 'antiword' не найдена. Используется fallback-метод (docx2txt).")

    # --- Метод 2: Fallback с использованием 'docx2txt' ---
    try:
        text = docx2txt.process(doc_path)
        if text and text.strip():
            print(f"INFO: Текст из {doc_path} извлечен с помощью docx2txt (fallback).")
            return text.strip()
        else:
            return f"ERROR: Fallback-метод (docx2txt) не смог извлечь текст из {doc_path}."
    except Exception as e:
        return f"ERROR: Не удалось обработать {doc_path}. Основной метод (antiword) и fallback (docx2txt) не сработали. Ошибка: {e}"

# --- ЗАДАЧА 3: ОБНОВЛЕННЫЙ РОУТЕР process_file ---
# --------------------------------------------------

def process_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    handlers = {
        '.txt': extract_text_from_txt,
        '.html': extract_text_from_txt, # Можно использовать один и тот же для простоты
        '.xml': extract_text_from_txt,
        '.docx': extract_text_from_docx,
        '.xlsx': extract_text_from_excel,
        '.xls': extract_text_from_excel,
        '.csv': extract_text_from_csv,
        '.png': extract_text_from_image,
        '.jpg': extract_text_from_image,
        '.jpeg': extract_text_from_image,
        '.pdf': extract_text_from_scanned_pdf,
        '.doc': extract_text_from_doc,
    }
    handler = handlers.get(ext)
    if handler:
        print(f"\n--- Processing {file_path} (handler: {handler.__name__}) ---")
        return handler(file_path)
    return f"ERROR: Нет обработчика для формата {ext}"

# --- ЗАДАЧА 4: СОЗДАНИЕ ТЕСТОВОЙ СРЕДЫ ---
# ------------------------------------------

def setup_test_environment():
    print("INFO: Создание тестовых файлов...")
    
    # CSV с русской кодировкой
    df_csv = pd.DataFrame([['Имя', 'Возраст'], ['Алексей', 30], ['Мария', 25]])
    df_csv.to_csv('test.csv', index=False, header=False, encoding='utf-8')

    # DOCX
    doc = docx.Document()
    doc.add_paragraph("Это тестовый документ DOCX.")
    doc.add_paragraph("Он содержит русский текст.")
    doc.save('test.docx')

    # XLSX
    df_excel = pd.DataFrame({'Продукт': ['Яблоки', 'Бананы'], 'Цена': [100, 120]})
    df_excel.to_excel('test.xlsx', index=False, sheet_name='Цены')

    # PNG с текстом
    img = Image.new('RGB', (400, 100), color = 'white')
    d = ImageDraw.Draw(img)
    try:
        # Попытка загрузить шрифт, если он есть
        font = ImageFont.truetype("arial.ttf", 20)
    except IOError:
        # Использовать стандартный шрифт, если нужный не найден
        font = ImageFont.load_default()
    d.text((10,10), "Тест OCR на изображении\nПроверка русского текста", fill='black', font=font)
    img.save('test.png')

    # Напоминание о файлах, которые нужно добавить вручную
    print("\n!!! ВНИМАНИЕ !!!")
    print("Пожалуйста, добавьте в эту же директорию:")
    print("1. Файл 'test.doc' (старый бинарный формат Word 97-2003).")
    print("2. Файл 'test_scanned.pdf' (отсканированный PDF без текстового слоя).")
    print("=========================================\n")


# --- ЗАДАЧА 5: ЗАПУСК ТЕСТОВ ---
# --------------------------------

if __name__ == "__main__":
    setup_test_environment()
    
    test_files = [
        'test.csv', 'test.docx', 'test.xlsx', 'test.png',
        # Добавьте сюда имена ваших файлов для ручного теста
        'test.doc', 'test_scanned.pdf' 
    ]
    
    for filename in test_files:
        if os.path.exists(filename):
            result = process_file(filename)
            print(f"Result: {result[:200]}..." if len(result) > 200 else f"Result: {result}")
        else:
            print(f"\n--- Processing {filename} ---")
            print(f"Result: SKIPPED (файл не найден).")
