#!/usr/bin/env python3
"""
Section-Aware Text Extraction for SupplierFinder Analysis
Supports PDF, DOC, DOCX, TXT, XLS, XLSX, and image formats
Enhanced with section detection and structure preservation
"""

import sys
import os
import tempfile
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    from docx import Document
    import pandas as pd
    import openpyxl
except ImportError as e:
    print(f"Missing required package: {e}", file=sys.stderr)
    sys.exit(1)

def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyPDF2 and OCR fallback with page references"""
    text = ""
    
    try:
        # First try PyPDF2 for text-based PDFs
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    # Add page marker for reference extraction
                    text += f"\n[PAGE:{page_num}]\n{page_text}\n"
        
        # If no text extracted, use OCR
        if not text.strip():
            images = convert_from_path(file_path)
            for page_num, image in enumerate(images, 1):
                ocr_text = pytesseract.image_to_string(image, lang='rus+eng')
                if ocr_text.strip():
                    text += f"\n[PAGE:{page_num}]\n{ocr_text}\n"
    
    except Exception as e:
        print(f"Error extracting from PDF: {e}", file=sys.stderr)
    
    return text

def extract_text_from_doc(file_path):
    """Extract text from DOC files using antiword"""
    import subprocess
    text = ""
    
    try:
        # Use antiword to convert DOC to text
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True, encoding='utf-8', errors='replace')
        if result.returncode == 0:
            raw_text = result.stdout
            # Add page markers - approximate based on content length
            lines = raw_text.split('\n')
            lines_per_page = 50  # Approximate lines per page
            page_count = 1
            line_count = 0
            
            text += f"[PAGE:{page_count}]\n"
            for line in lines:
                if line.strip():  # Only process non-empty lines
                    text += line + "\n"
                    line_count += 1
                    
                    # Create logical page breaks
                    if line_count >= lines_per_page:
                        page_count += 1
                        text += f"\n[PAGE:{page_count}]\n"
                        line_count = 0
        else:
            print(f"antiword failed with exit code {result.returncode}: {result.stderr}", file=sys.stderr)
            
    except Exception as e:
        print(f"Error extracting from DOC: {e}", file=sys.stderr)
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX files with enhanced processing and comprehensive content extraction"""
    text = ""
    
    try:
        # Primary method: python-docx with complete content extraction
        doc = Document(file_path)
        paragraph_count = 0
        page_count = 1
        paragraphs_per_page = 30  # Adjusted for better coverage
        
        text += f"[PAGE:{page_count}]\n"
        
        # Extract ALL paragraphs with enhanced processing
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                # Preserve paragraph structure and numbering
                text += para_text + "\n"
                paragraph_count += 1
                
                # Create logical page breaks based on content density
                if paragraph_count >= paragraphs_per_page:
                    page_count += 1
                    text += f"\n[PAGE:{page_count}]\n"
                    paragraph_count = 0
        
        # Extract tables content
        for table in doc.tables:
            text += "\n=== ТАБЛИЦА ===\n"
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    text += " | ".join(row_cells) + "\n"
            text += "=== КОНЕЦ ТАБЛИЦЫ ===\n\n"
        
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                header_text = ""
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        header_text += paragraph.text.strip() + " "
                if header_text:
                    text += f"\n=== ЗАГОЛОВОК: {header_text.strip()} ===\n"
            
            if section.footer:
                footer_text = ""
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        footer_text += paragraph.text.strip() + " "
                if footer_text:
                    text += f"\n=== ПОДВАЛ: {footer_text.strip()} ===\n"
                    
    except Exception as e:
        print(f"Primary DOCX extraction failed: {e}", file=sys.stderr)
        # Fallback method: try with zipfile approach
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            text = extract_docx_with_zipfile(file_path)
            if not text.strip():
                # Second fallback: try with subprocess and pandoc
                text = extract_docx_with_pandoc(file_path)
                
        except Exception as fallback_error:
            print(f"DOCX fallback extraction failed: {fallback_error}", file=sys.stderr)
    
    return text

def extract_docx_with_zipfile(file_path):
    """Fallback DOCX extraction using zipfile and XML parsing"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    text = ""
    try:
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # Extract main document content
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # Define namespace
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            page_count = 1
            text += f"[PAGE:{page_count}]\n"
            
            # Extract text from all text elements
            for elem in root.iter():
                if elem.tag.endswith('}t'):  # Text elements
                    if elem.text:
                        text += elem.text
                elif elem.tag.endswith('}p'):  # Paragraph breaks
                    text += "\n"
                elif elem.tag.endswith('}br'):  # Line breaks
                    if elem.get('type') == 'page':
                        page_count += 1
                        text += f"\n[PAGE:{page_count}]\n"
                    else:
                        text += "\n"
                        
    except Exception as e:
        print(f"Zipfile DOCX extraction failed: {e}", file=sys.stderr)
        
    return text

def extract_docx_with_pandoc(file_path):
    """Third fallback: use pandoc if available"""
    import subprocess
    
    text = ""
    try:
        # Try pandoc conversion
        result = subprocess.run([
            'pandoc', file_path, '-t', 'plain', '--wrap=none'
        ], capture_output=True, text=True, encoding='utf-8', errors='replace')
        
        if result.returncode == 0:
            raw_text = result.stdout
            # Add page markers based on content length
            lines = raw_text.split('\n')
            lines_per_page = 40
            page_count = 1
            line_count = 0
            
            text += f"[PAGE:{page_count}]\n"
            for line in lines:
                if line.strip():
                    text += line + "\n"
                    line_count += 1
                    
                    if line_count >= lines_per_page:
                        page_count += 1
                        text += f"\n[PAGE:{page_count}]\n"
                        line_count = 0
        else:
            print(f"Pandoc failed: {result.stderr}", file=sys.stderr)
            
    except Exception as e:
        print(f"Pandoc extraction failed: {e}", file=sys.stderr)
        
    return text

def extract_text_from_excel(file_path):
    """Extract text from Excel files (XLS, XLSX)"""
    text = ""
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            # Convert dataframe to text representation
            sheet_text = f"=== Sheet: {sheet_name} ===\n"
            sheet_text += df.to_string(index=False, na_rep='') + "\n\n"
            text += sheet_text
    except Exception as e:
        print(f"Error extracting from Excel: {e}", file=sys.stderr)
    
    return text

def extract_text_from_image(file_path):
    """Extract text from image files using OCR"""
    text = ""
    
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='rus+eng')
    except Exception as e:
        print(f"Error extracting from image: {e}", file=sys.stderr)
    
    return text

def extract_text_from_txt(file_path):
    """Extract text from plain text files with robust encoding detection"""
    text = ""
    
    try:
        # Try different encodings in order of preference
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'iso-8859-1', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    text = file.read()
                    # If we got meaningful text, break
                    if len(text.strip()) > 0:
                        break
            except (UnicodeDecodeError, UnicodeError):
                continue
                
        # If still no text, try binary read with error handling
        if not text.strip():
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                # Try to decode with error replacement
                for encoding in encodings:
                    try:
                        text = raw_data.decode(encoding, errors='replace')
                        if len(text.strip()) > 0:
                            break
                    except:
                        continue
                        
    except Exception as e:
        print(f"Error extracting from TXT: {e}", file=sys.stderr)
    
    return text

def detect_file_type(file_path):
    """Detect actual file type using magic numbers and file command"""
    import subprocess
    
    try:
        # Use file command to detect actual file type
        result = subprocess.run(['file', '-b', '--mime-type', file_path], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            detected_mime = result.stdout.strip()
            print(f"Detected MIME type: {detected_mime}", file=sys.stderr)
            return detected_mime
    except Exception as e:
        print(f"File type detection failed: {e}", file=sys.stderr)
    
    # Fallback: check magic numbers manually
    try:
        with open(file_path, 'rb') as f:
            header = f.read(16)
            
        # PDF files start with %PDF
        if header.startswith(b'%PDF'):
            return 'application/pdf'
        # ZIP-based files (DOCX, XLSX)
        elif header.startswith(b'PK\x03\x04'):
            # Check for specific Office formats
            try:
                import zipfile
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    files = zip_file.namelist()
                    if 'word/document.xml' in files:
                        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    elif 'xl/workbook.xml' in files:
                        return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    else:
                        return 'application/zip'
            except:
                return 'application/zip'
        # Old Office formats
        elif header.startswith(b'\xd0\xcf\x11\xe0'):
            return 'application/msword'  # Could also be Excel or PowerPoint
        # JPEG
        elif header.startswith(b'\xff\xd8\xff'):
            return 'image/jpeg'
        # PNG
        elif header.startswith(b'\x89PNG'):
            return 'image/png'
        # Text files (UTF-8 BOM)
        elif header.startswith(b'\xef\xbb\xbf'):
            return 'text/plain'
            
    except Exception as e:
        print(f"Magic number detection failed: {e}", file=sys.stderr)
    
    return None

def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_text.py <file_path> <mime_type>", file=sys.stderr)
        sys.exit(1)
    
    file_path = sys.argv[1]
    provided_mime_type = sys.argv[2]
    
    # Log file information for debugging
    print(f"Processing file: {file_path}", file=sys.stderr)
    print(f"Provided MIME type: {provided_mime_type}", file=sys.stderr)
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    # Get file size for debugging
    try:
        file_size = os.path.getsize(file_path)
        print(f"File size: {file_size} bytes", file=sys.stderr)
    except Exception as e:
        print(f"Error getting file size: {e}", file=sys.stderr)
    
    # Detect actual file type
    detected_mime_type = detect_file_type(file_path)
    
    # Use detected type if available, otherwise fall back to provided type
    mime_type = detected_mime_type if detected_mime_type else provided_mime_type
    print(f"Using MIME type: {mime_type}", file=sys.stderr)
    
    text = ""
    file_ext = Path(file_path).suffix.lower()
    
    # Route to appropriate extraction function based on file type
    try:
        if mime_type == 'application/pdf' or file_ext == '.pdf':
            print("Using PDF extraction", file=sys.stderr)
            text = extract_text_from_pdf(file_path)
        elif mime_type == 'application/msword' or file_ext == '.doc':
            print("Using DOC extraction", file=sys.stderr)
            text = extract_text_from_doc(file_path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == '.docx':
            print("Using DOCX extraction", file=sys.stderr)
            text = extract_text_from_docx(file_path)
        elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'] or file_ext in ['.xls', '.xlsx']:
            print("Using Excel extraction", file=sys.stderr)
            text = extract_text_from_excel(file_path)
        elif mime_type.startswith('image/') or file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
            print("Using image OCR extraction", file=sys.stderr)
            text = extract_text_from_image(file_path)
        elif mime_type == 'text/plain' or file_ext == '.txt':
            print("Using text extraction", file=sys.stderr)
            text = extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file type: {mime_type}, extension: {file_ext}", file=sys.stderr)
            sys.exit(1)
            
    except Exception as e:
        print(f"Critical error during text extraction: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        text = ""
    
    # Clean and output the text
    text = text.strip()
    extracted_length = len(text)
    print(f"Extracted text length: {extracted_length} characters", file=sys.stderr)
    
    if text:
        # Ensure UTF-8 output
        try:
            # Print with explicit UTF-8 encoding
            sys.stdout.buffer.write(text.encode('utf-8'))
            sys.stdout.buffer.flush()
        except Exception as e:
            print(f"Error outputting text: {e}", file=sys.stderr)
            # Fallback to regular print with error replacement
            print(text.encode('utf-8', errors='replace').decode('utf-8'))
    else:
        print(f"No text extracted from file", file=sys.stderr)
        # Don't exit with error - just output empty result
        print("")

if __name__ == "__main__":
    main()