#!/usr/bin/env python3
"""
Тестирует разные методы извлечения текста из DOC файлов
"""

import sys
import os
import tempfile
import docx2txt
import docx
from pathlib import Path

def test_docx2txt(file_path):
    """Тестирует docx2txt"""
    try:
        print("Testing docx2txt...")
        text = docx2txt.process(file_path)
        if text and len(text.strip()) > 10:
            print(f"SUCCESS docx2txt: {len(text)} chars")
            print(f"Preview: {text[:200]}...")
            return text.strip()
        else:
            print("FAILED docx2txt: Empty or very short text")
            return None
    except Exception as e:
        print(f"FAILED docx2txt: {e}")
        return None

def test_python_docx(file_path):
    """Тестирует python-docx"""
    try:
        print("Testing python-docx...")
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        text = '\n'.join(full_text)
        if text and len(text.strip()) > 10:
            print(f"SUCCESS python-docx: {len(text)} chars")
            print(f"Preview: {text[:200]}...")
            return text.strip()
        else:
            print("FAILED python-docx: Empty or very short text")
            return None
    except Exception as e:
        print(f"FAILED python-docx: {e}")
        return None

def test_binary_extraction(file_path):
    """Тестирует бинарное извлечение (текущий метод)"""
    try:
        print("Testing binary extraction...")
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # Try to decode as UTF-8 with error handling
        import re
        text_parts = []
        
        try:
            decoded = content.decode('utf-8', errors='ignore')
            readable_text = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()\-]+', decoded)
            if readable_text:
                text_parts.extend(readable_text)
        except:
            pass
        
        if text_parts:
            combined_text = ' '.join(text_parts)
            cleaned_text = re.sub(r'\s+', ' ', combined_text).strip()
            if cleaned_text and len(cleaned_text) > 10:
                print(f"SUCCESS Binary extraction: {len(cleaned_text)} chars")
                print(f"Preview: {cleaned_text[:200]}...")
                return cleaned_text
        
        print("FAILED Binary extraction: Empty or very short text")
        return None
    except Exception as e:
        print(f"FAILED Binary extraction: {e}")
        return None

def main():
    if len(sys.argv) != 2:
        print("Usage: python test-doc-extraction-methods.py <doc_file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}")
        sys.exit(1)
    
    print(f"Testing DOC extraction methods for: {file_path}")
    print("=" * 60)
    
    # Test different methods
    methods = [
        ("docx2txt", test_docx2txt),
        ("python-docx", test_python_docx),
        ("binary extraction", test_binary_extraction)
    ]
    
    results = {}
    
    for method_name, method_func in methods:
        print(f"\nMethod: {method_name}")
        print("-" * 40)
        result = method_func(file_path)
        results[method_name] = result
        print()
    
    # Summary
    print("SUMMARY")
    print("=" * 60)
    
    best_method = None
    best_length = 0
    
    for method_name, result in results.items():
        if result:
            length = len(result)
            print(f"SUCCESS {method_name}: {length} chars")
            if length > best_length:
                best_length = length
                best_method = method_name
        else:
            print(f"FAILED {method_name}: Failed")
    
    if best_method:
        print(f"\nBest method: {best_method} ({best_length} chars)")
        print(f"Best result preview: {results[best_method][:300]}...")
    else:
        print("\nFAILED: All methods failed")

if __name__ == "__main__":
    main()
   