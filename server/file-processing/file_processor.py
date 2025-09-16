#!/usr/bin/env python3
"""
File Processing Module for SupplierFinder

This module extracts text from various file attachments in supplier emails including:
- PDF (text and scanned)
- Microsoft Word documents (DOC/DOCX)
- Microsoft Excel spreadsheets (XLS/XLSX)
- Images (JPG, PNG, TIFF)

The extracted text is stored in structured formats (JSON/CSV) for further analysis.
"""

import os
import sys
import json
import csv
import logging
import re
from datetime import datetime
import traceback
from typing import Dict, List, Any, Tuple, Optional, Union
import tempfile
import shutil

# PDF processing
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdf2image import convert_from_path, convert_from_bytes

# Document processing
import docx
import docx2txt
import subprocess

# Image processing
import pytesseract
from PIL import Image
import cv2

# Spreadsheet processing
import pandas as pd
import openpyxl
import xlrd

# File type detection
import magic

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("file_processing.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("file_processor")


class FileProcessor:
    """Main class for processing file attachments"""
    
    def __init__(self, output_dir: str = "processed_attachments"):
        """
        Initialize the FileProcessor
        
        Args:
            output_dir: Directory to store processed output files
        """
        self.output_dir = output_dir
        self.error_dir = os.path.join(output_dir, "errors")
        
        # Create output directories if they don't exist
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.error_dir, exist_ok=True)
        
        # Tesseract configuration
        self.tesseract_config = "--oem 3 --psm 6"
        
        logger.info(f"FileProcessor initialized with output directory: {output_dir}")
    
    def process_attachment(self, 
                          file_data: bytes, 
                          filename: str, 
                          supplier_name: str, 
                          supplier_email: str, 
                          email_date: str,
                          request_id: Optional[int] = None) -> Dict[str, Any]:
        """
        Process a single attachment file
        
        Args:
            file_data: Binary content of the file
            filename: Name of the file
            supplier_name: Name of the supplier
            supplier_email: Email of the supplier
            email_date: Date the email was received
            request_id: Optional ID of the request
            
        Returns:
            Dictionary with extracted data and metadata
        """
        logger.info(f"Processing attachment: {filename} from {supplier_name} ({supplier_email})")
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        # Create a temporary file to work with
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            # Use new fallback processing system
            try:
                processed_result, processing_status = self._process_with_fallback(temp_file_path, filename)
                
                # Handle result based on type
                if isinstance(processed_result, dict):
                    # Excel result
                    extracted_data = processed_result
                    extracted_text = self._excel_data_to_text(processed_result)
                else:
                    # Text result
                    extracted_text = processed_result
                    extracted_data = None
                    
            except Exception as e:
                logger.error(f"Fallback processing failed for {filename}: {e}")
                processing_status = self._create_user_friendly_error(e, filename)
                extracted_text = f"Ошибка обработки: {processing_status['user_message']}"
                extracted_data = None
            
            result = {
                "metadata": {
                    "supplier_name": supplier_name,
                    "supplier_email": supplier_email,
                    "email_date": email_date,
                    "filename": filename,
                    "file_type": file_ext.lstrip('.'),
                    "processing_date": datetime.now().isoformat(),
                    "request_id": request_id
                },
                "text_content": extracted_text[:10000] if len(extracted_text) > 10000 else extracted_text,
                "data": extracted_data,
                "processing_status": processing_status
            }
            
            # Save the result
            self._save_processing_result(result, filename, supplier_email)
            
            return result
        
        except Exception as e:
            error_msg = f"Error processing {filename}: {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            
            # Save error information
            error_data = {
                "error": str(e),
                "traceback": traceback.format_exc(),
                "file": filename,
                "supplier": supplier_name,
                "email": supplier_email
            }
            
            error_filename = f"error_{os.path.basename(filename)}_{datetime.now().strftime('%Y%m%d%H%M%S')}.json"
            with open(os.path.join(self.error_dir, error_filename), 'w', encoding='utf-8') as f:
                json.dump(error_data, f, ensure_ascii=False, indent=2)
            
            return {
                "metadata": {
                    "supplier_name": supplier_name,
                    "supplier_email": supplier_email,
                    "email_date": email_date,
                    "filename": filename,
                    "file_type": file_ext.lstrip('.'),
                    "processing_date": datetime.now().isoformat(),
                    "request_id": request_id,
                    "error": True,
                    "error_message": str(e)
                },
                "text_content": "Error during processing",
                "data": None
            }
        
        finally:
            # Clean up temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for maximum reliability"""
        logger.info(f"Processing PDF: {file_path}")
        print(f"[PDF] Processing PDF file: {file_path}")  # Print to stdout for attachment analyzer
        
        all_text = []
        errors = []
        is_scanned_pdf = False  # Flag to track if this is likely a scanned PDF
        
        # 1. Try PyPDF2 first (fastest method)
        try:
            logger.info("Attempting extraction with PyPDF2")
            print("[PDF] Attempting extraction with PyPDF2")
            pypdf_text = ""
            page_count = 0
            content_pages = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning(f"Encrypted PDF detected: {file_path}")
                    print(f"[PDF] WARNING: Encrypted PDF detected: {file_path}")
                    errors.append("PDF is encrypted")
                else:
                    # Get number of pages
                    num_pages = len(pdf_reader.pages)
                    logger.info(f"PDF has {num_pages} pages")
                    print(f"[PDF] PDF has {num_pages} pages")
                    
                    # Extract text from each page
                    for page_num in range(num_pages):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            pypdf_text += page_text + "\n\n"
                            content_pages += 1
                        
                        page_count += 1
                            
                        # Print debug info periodically
                        if page_num % 10 == 0 or page_num == num_pages - 1:
                            print(f"[PDF] Processed page {page_num+1}/{num_pages} with PyPDF2")
                    
                    # Check if this is likely a scanned PDF
                    # If less than 30% of pages contain extractable text, it's likely a scanned document
                    if page_count > 0 and content_pages / page_count < 0.3:
                        is_scanned_pdf = True
                        print(f"[PDF] This appears to be a scanned PDF: {content_pages}/{page_count} pages have extractable text")
            
            # Check if we got any text
            if pypdf_text.strip():
                text_length = len(pypdf_text)
                logger.info(f"Successfully extracted {text_length} chars from {content_pages}/{page_count} pages with PyPDF2")
                print(f"[PDF] Successfully extracted {text_length} chars from {content_pages}/{page_count} pages with PyPDF2")
                
                # Print a sample of the text
                sample_length = min(200, text_length)
                print(f"[PDF] Sample PyPDF2 text: {pypdf_text[:sample_length]}...")
                
                all_text.append(pypdf_text)
            else:
                logger.warning("PyPDF2 extraction returned empty or whitespace-only text")
                print("[PDF] WARNING: PyPDF2 extraction returned empty or whitespace-only text")
                is_scanned_pdf = True  # If no text was extracted, it's likely a scanned document
        except Exception as e:
            logger.error(f"PyPDF2 extraction error: {str(e)}")
            print(f"[PDF] ERROR: PyPDF2 extraction failed: {str(e)}")
            traceback_str = traceback.format_exc()
            logger.error(traceback_str)
            print(f"[PDF] Traceback: {traceback_str}")
            errors.append(f"PyPDF2: {str(e)}")
        
        # 2. Try PDFMiner (more robust but slower)
        try:
            logger.info("Attempting extraction with PDFMiner")
            print("[PDF] Attempting extraction with PDFMiner")
            
            pdfminer_text = pdfminer_extract_text(file_path)
            
            if pdfminer_text and pdfminer_text.strip():
                text_length = len(pdfminer_text)
                logger.info(f"Successfully extracted {text_length} chars with PDFMiner")
                print(f"[PDF] Successfully extracted {text_length} chars with PDFMiner")
                
                # Print a sample of the text
                sample_length = min(200, text_length)
                print(f"[PDF] Sample PDFMiner text: {pdfminer_text[:sample_length]}...")
                
                all_text.append(pdfminer_text)
                
                # If we got substantial text with PDFMiner but not with PyPDF2, 
                # it's probably not a scanned PDF but a complex PDF
                if text_length > 500 and is_scanned_pdf:
                    is_scanned_pdf = False
            else:
                logger.warning("PDFMiner extraction returned empty or whitespace-only text")
                print("[PDF] WARNING: PDFMiner extraction returned empty or whitespace-only text")
                
                # If both PyPDF2 and PDFMiner failed, it's very likely a scanned PDF
                if not all_text:
                    is_scanned_pdf = True
        except Exception as e:
            logger.error(f"PDFMiner extraction error: {str(e)}")
            print(f"[PDF] ERROR: PDFMiner extraction failed: {str(e)}")
            traceback_str = traceback.format_exc()
            logger.error(traceback_str)
            print(f"[PDF] Traceback: {traceback_str}")
            errors.append(f"PDFMiner: {str(e)}")
            
            # If both extraction methods failed, it's likely a scanned PDF or corrupted
            if not all_text:
                is_scanned_pdf = True
        
        # 3. For scanned PDFs or if textual extraction yielded very little, use OCR
        if is_scanned_pdf or not all_text or max(len(t) for t in all_text if t) < 500:
            try:
                if is_scanned_pdf:
                    logger.info("Detected scanned PDF, performing OCR")
                    print("[PDF] Detected scanned PDF, performing OCR")
                else:
                    logger.info("Text extraction produced limited results, attempting OCR")
                    print("[PDF] Text extraction produced limited results, attempting OCR")
                
                ocr_text = self._ocr_pdf(file_path)
                
                if ocr_text and ocr_text.strip():
                    text_length = len(ocr_text)
                    logger.info(f"Successfully extracted {text_length} chars with OCR")
                    print(f"[PDF] Successfully extracted {text_length} chars with OCR")
                    
                    # Print a sample of the text
                    sample_length = min(200, text_length)
                    print(f"[PDF] Sample OCR text: {ocr_text[:sample_length]}...")
                    
                    all_text.append(ocr_text)
                else:
                    logger.warning("OCR extraction returned empty or whitespace-only text")
                    print("[PDF] WARNING: OCR extraction returned empty or whitespace-only text")
            except Exception as e:
                logger.error(f"OCR extraction error: {str(e)}")
                print(f"[PDF] ERROR: OCR extraction failed: {str(e)}")
                traceback_str = traceback.format_exc()
                logger.error(traceback_str)
                print(f"[PDF] Traceback: {traceback_str}")
                errors.append(f"OCR: {str(e)}")
        
        # Combine all extracted text, prioritizing longer results
        if all_text:
            # Filter out empty texts and sort by length (longest first)
            valid_texts = [t for t in all_text if t and t.strip()]
            valid_texts.sort(key=len, reverse=True)
            
            if valid_texts:
                # Use the longest text as the base
                longest_text = valid_texts[0]
                
                # If the longest text is very short but we have multiple sources,
                # combine them all for maximum content
                if len(longest_text) < 1000 and len(valid_texts) > 1:
                    combined_text = "\n\n===\n\n".join(valid_texts)
                    text_length = len(combined_text)
                    logger.info(f"Combined all extracted texts: {text_length} chars")
                    print(f"[PDF] Combined all extracted texts: {text_length} chars from {len(valid_texts)} sources")
                else:
                    # Otherwise just use the longest text
                    text_length = len(longest_text)
                    combined_text = longest_text
                    logger.info(f"Using longest extracted text: {text_length} chars")
                    print(f"[PDF] Using longest extracted text: {text_length} chars")
                
                # Print final text statistics
                print(f"[PDF] Final extracted text length: {len(combined_text)} characters")
                sample_length = min(300, len(combined_text))
                print(f"[PDF] Final text sample: {combined_text[:sample_length]}...")
                
                return combined_text
        
        # If we reach here, all extraction methods failed
        error_msg = f"All text extraction methods failed: {'; '.join(errors)}"
        logger.error(error_msg)
        print(f"[PDF] ERROR: {error_msg}")
        return f"Text extraction failed with errors: {'; '.join(errors)}"
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Process a PDF using OCR with enhanced image preprocessing for better quality results"""
        logger.info(f"Performing OCR on PDF: {file_path}")
        print(f"[OCR] Performing OCR on PDF: {file_path}")
        
        try:
            # Create a temporary directory for images
            temp_dir = tempfile.mkdtemp()
            print(f"[OCR] Created temporary directory: {temp_dir}")
            
            # Convert PDF to images with higher DPI for better quality
            print(f"[OCR] Converting PDF to images at 300 DPI...")
            images = convert_from_path(file_path, dpi=300)  # Higher DPI for better quality
            num_pages = len(images)
            print(f"[OCR] PDF converted to {num_pages} images")
            
            # Limit to first 20 pages to avoid excessive processing
            max_pages = min(20, num_pages)
            if max_pages < num_pages:
                print(f"[OCR] Processing only first {max_pages} pages out of {num_pages} total")
            
            text = ""
            processed_pages = 0
            
            for i, image in enumerate(images):
                if i >= max_pages:
                    print(f"[OCR] Reached maximum page limit ({max_pages})")
                    break
                
                # Save the image temporarily
                temp_img_path = os.path.join(temp_dir, f"page_{i}.png")
                print(f"[OCR] Saving page {i+1}/{num_pages} to {temp_img_path}")
                image.save(temp_img_path, 'PNG')
                
                try:
                    # Process with multiple preprocessing techniques for better results
                    print(f"[OCR] Preprocessing image for page {i+1} with multiple techniques")
                    img = cv2.imread(temp_img_path)
                    
                    if img is None:
                        print(f"[OCR] WARNING: Could not load image with OpenCV: {temp_img_path}")
                        # Try direct OCR without preprocessing
                        print(f"[OCR] Attempting direct OCR without preprocessing")
                        page_text = pytesseract.image_to_string(Image.open(temp_img_path), config=self.tesseract_config)
                    else:
                        # Try multiple preprocessing techniques and use the best result
                        preprocessing_results = []
                        
                        # 1. Original image without preprocessing (baseline)
                        print(f"[OCR] Method 1: Original image without preprocessing")
                        original_text = pytesseract.image_to_string(Image.open(temp_img_path), config=self.tesseract_config)
                        preprocessing_results.append((original_text, "original"))
                        
                        # 2. Basic grayscale + threshold
                        print(f"[OCR] Method 2: Basic grayscale + threshold")
                        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                        _, threshold = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        threshold_path = os.path.join(temp_dir, f"page_{i}_threshold.png")
                        cv2.imwrite(threshold_path, threshold)
                        threshold_text = pytesseract.image_to_string(Image.open(threshold_path), config=self.tesseract_config)
                        preprocessing_results.append((threshold_text, "threshold"))
                        
                        # 3. Enhanced contrast
                        print(f"[OCR] Method 3: Enhanced contrast")
                        # Apply adaptive equalization to enhance contrast
                        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                        contrast_enhanced = clahe.apply(gray)
                        # Apply threshold to enhanced image
                        _, enhanced_threshold = cv2.threshold(contrast_enhanced, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        enhanced_path = os.path.join(temp_dir, f"page_{i}_enhanced.png")
                        cv2.imwrite(enhanced_path, enhanced_threshold)
                        enhanced_text = pytesseract.image_to_string(Image.open(enhanced_path), config=self.tesseract_config)
                        preprocessing_results.append((enhanced_text, "enhanced"))
                        
                        # 4. Noise reduction
                        print(f"[OCR] Method 4: Noise reduction")
                        # Apply denoising 
                        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
                        _, denoised_threshold = cv2.threshold(denoised, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        denoised_path = os.path.join(temp_dir, f"page_{i}_denoised.png")
                        cv2.imwrite(denoised_path, denoised_threshold)
                        denoised_text = pytesseract.image_to_string(Image.open(denoised_path), config=self.tesseract_config)
                        preprocessing_results.append((denoised_text, "denoised"))
                        
                        # 5. Advanced: RGB channels separation
                        # This can help with colored backgrounds and watermarks
                        print(f"[OCR] Method 5: RGB channel separation")
                        b, g, r = cv2.split(img)
                        # Use green channel which typically has best contrast for text
                        _, green_threshold = cv2.threshold(g, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        green_path = os.path.join(temp_dir, f"page_{i}_green.png")
                        cv2.imwrite(green_path, green_threshold)
                        green_text = pytesseract.image_to_string(Image.open(green_path), config=self.tesseract_config)
                        preprocessing_results.append((green_text, "green_channel"))
                        
                        # Compare results and choose the best one (most text content)
                        valid_results = [(text, method) for text, method in preprocessing_results if text and text.strip()]
                        if valid_results:
                            # Sort by text length and choose the longest result
                            valid_results.sort(key=lambda x: len(x[0]), reverse=True)
                            page_text, best_method = valid_results[0]
                            print(f"[OCR] Best result from method '{best_method}' with {len(page_text)} chars")
                        else:
                            print(f"[OCR] All preprocessing methods failed to extract text")
                            page_text = ""
                        
                    # Add text to result
                    if page_text and page_text.strip():
                        print(f"[OCR] Extracted {len(page_text)} chars from page {i+1}")
                        text += page_text + "\n\n"
                        processed_pages += 1
                        
                        # Print sample of extracted text
                        sample_length = min(100, len(page_text))
                        print(f"[OCR] Sample text from page {i+1}: {page_text[:sample_length]}...")
                    else:
                        print(f"[OCR] No text extracted from page {i+1}")
                    
                except Exception as page_error:
                    print(f"[OCR] Error processing page {i+1}: {str(page_error)}")
                    logger.error(f"Error processing page {i+1}: {str(page_error)}")
                    logger.error(traceback.format_exc())
            
            # Clean up temporary directory
            print(f"[OCR] Cleaning up temporary directory: {temp_dir}")
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            # Post-process text to improve quality
            if text and text.strip():
                # Clean up text by removing multiple newlines and spaces
                cleaned_text = re.sub(r'\n{3,}', '\n\n', text)  # Replace 3+ newlines with 2
                cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)  # Replace multiple spaces with one
                
                print(f"[OCR] Successfully extracted {len(cleaned_text)} chars from {processed_pages} pages")
                return cleaned_text
            else:
                print(f"[OCR] No text extracted from any of the {processed_pages} processed pages")
                return ""
                
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            print(f"[OCR] ERROR: OCR process failed: {str(e)}")
            traceback_str = traceback.format_exc()
            logger.error(traceback_str)
            print(f"[OCR] Traceback: {traceback_str}")
            
            # Clean up any temporary directory
            try:
                temp_dir_to_clean = locals().get('temp_dir')
                if temp_dir_to_clean and os.path.exists(temp_dir_to_clean):
                    print(f"[OCR] Cleaning up temporary directory after error: {temp_dir_to_clean}")
                    shutil.rmtree(temp_dir_to_clean, ignore_errors=True)
            except Exception as cleanup_error:
                print(f"[OCR] Error during cleanup: {str(cleanup_error)}")
                
            raise
    
    def _process_word(self, file_path: str) -> str:
        """Extract text from Word document"""
        logger.info(f"Processing Word document: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error in Word processing: {str(e)}")
            raise

    def _process_old_word(self, file_path: str) -> str:
        """
        Process older Word document (.doc) using multiple fallback methods
        """
        logger.info(f"Processing old Word document (.doc): {file_path}")
        
        # Method 1: Try docx2txt first - it can handle some .doc files
        try:
            logger.info("Trying docx2txt for .doc file")
            text = docx2txt.process(file_path)
            if text and len(text.strip()) > 10:
                logger.info(f"Successfully extracted {len(text)} characters using docx2txt")
                return text.strip()
            else:
                logger.warning("docx2txt returned empty or very short text")
        except Exception as e:
            logger.warning(f"docx2txt failed: {str(e)}")
        
        # Method 2: Try reading as binary and extracting text manually
        try:
            logger.info("Trying binary text extraction for .doc file")
            text = self._extract_text_from_binary_doc(file_path)
            if text and len(text.strip()) > 10:
                logger.info(f"Successfully extracted {len(text)} characters using binary extraction")
                return text.strip()
        except Exception as e:
            logger.warning(f"Binary extraction failed: {str(e)}")
        
        # Method 3: Try antiword as fallback (if available)
        try:
            logger.info("Trying antiword for .doc file")
            result = subprocess.run(['antiword', file_path], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                logger.info(f"Successfully extracted {len(text)} characters using antiword")
                return text
            else:
                logger.warning(f"antiword returned error code {result.returncode}")
        except subprocess.TimeoutExpired:
            logger.warning("antiword timed out after 30 seconds")
        except Exception as e:
            logger.warning(f"antiword failed: {str(e)}")
        
        # Method 4: Try LibreOffice conversion (if available)
        try:
            logger.info("Trying LibreOffice conversion for .doc file")
            return self._process_doc_with_libreoffice(file_path)
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {str(e)}")
        
        # If all methods fail, raise error
        raise Exception("Failed to process .doc file with all available methods")
    
    def _extract_text_from_binary_doc(self, file_path: str) -> str:
        """
        Extract text from DOC file by reading binary content and looking for text patterns
        """
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Look for text patterns in the binary content
            # DOC files often contain readable text mixed with binary data
            text_parts = []
            
            # Try to decode as UTF-8 with error handling
            try:
                decoded = content.decode('utf-8', errors='ignore')
                # Extract readable text (letters, numbers, spaces, punctuation)
                import re
                readable_text = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()\-]+', decoded)
                if readable_text:
                    text_parts.extend(readable_text)
            except:
                pass
            
            # Try to decode as Windows-1251 (common for Russian text)
            try:
                decoded = content.decode('windows-1251', errors='ignore')
                import re
                readable_text = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()\-]+', decoded)
                if readable_text:
                    text_parts.extend(readable_text)
            except:
                pass
            
            # Try to decode as Latin-1
            try:
                decoded = content.decode('latin-1', errors='ignore')
                import re
                readable_text = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()\-]+', decoded)
                if readable_text:
                    text_parts.extend(readable_text)
            except:
                pass
            
            if text_parts:
                # Join all text parts and clean up
                combined_text = ' '.join(text_parts)
                # Remove excessive whitespace
                import re
                cleaned_text = re.sub(r'\s+', ' ', combined_text).strip()
                return cleaned_text
            
            return ""
            
        except Exception as e:
            logger.warning(f"Binary text extraction failed: {str(e)}")
            return ""

    def _process_doc_with_libreoffice(self, file_path: str) -> str:
        """
        Convert DOC to DOCX using LibreOffice, then process with python-docx
        """
        logger.info(f"Converting DOC to DOCX using LibreOffice: {file_path}")
        
        # Create temporary directory for conversion
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Use LibreOffice to convert DOC to DOCX
            logger.info("Running LibreOffice conversion...")
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                '--outdir', temp_dir, file_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0:
                logger.warning(f"LibreOffice conversion failed with code {result.returncode}")
                logger.warning(f"LibreOffice stderr: {result.stderr}")
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
            # Find the converted DOCX file
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            converted_file = os.path.join(temp_dir, f"{base_name}.docx")
            
            if not os.path.exists(converted_file):
                raise Exception(f"Converted DOCX file not found: {converted_file}")
            
            logger.info(f"✅ Successfully converted to DOCX: {converted_file}")
            
            # Now process the DOCX file with python-docx
            doc = docx.Document(converted_file)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            result_text = "\n".join(full_text)
            logger.info(f"✅ Successfully extracted {len(result_text)} characters from converted DOCX")
            return result_text
            
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice conversion timed out after 60 seconds")
            raise Exception("LibreOffice conversion timed out")
        except Exception as e:
            logger.error(f"LibreOffice conversion failed: {str(e)}")
            raise
        finally:
            # Clean up temporary directory
            try:
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temp directory: {cleanup_error}")
    
    def _process_excel(self, file_path: str) -> Dict[str, List[List[Any]]]:
        """Extract data from Excel spreadsheet"""
        logger.info(f"Processing Excel spreadsheet: {file_path}")
        
        try:
            # Determine the engine to use based on file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.xls':
                # Use xlrd for older Excel formats
                engine = 'xlrd'
            else:
                # Use openpyxl for newer Excel formats
                engine = 'openpyxl'
            
            # Read all sheets
            excel_data = {}
            xls = pd.ExcelFile(file_path, engine=engine)
            
            for sheet_name in xls.sheet_names:
                # Read sheet into a DataFrame
                df = pd.read_excel(xls, sheet_name=sheet_name)
                
                # Convert DataFrame to a list of lists
                headers = df.columns.tolist()
                rows = df.values.tolist()
                
                excel_data[sheet_name] = {
                    "headers": headers,
                    "data": rows
                }
            
            return excel_data
        except Exception as e:
            logger.error(f"Error in Excel processing: {str(e)}")
            raise
    
    def _excel_data_to_text(self, excel_data: Dict[str, Any]) -> str:
        """Convert Excel data to text format for summary view"""
        text_output = []
        
        for sheet_name, sheet_data in excel_data.items():
            text_output.append(f"Sheet: {sheet_name}")
            text_output.append("-" * 50)
            
            # Add headers
            if sheet_data["headers"]:
                text_output.append(" | ".join(str(h) for h in sheet_data["headers"]))
                text_output.append("-" * 50)
            
            # Add data rows (limit to 100 rows for text summary)
            for i, row in enumerate(sheet_data["data"]):
                if i >= 100:
                    text_output.append(f"... and {len(sheet_data['data']) - 100} more rows")
                    break
                text_output.append(" | ".join(str(cell) for cell in row))
            
            text_output.append("\n")
        
        return "\n".join(text_output)
    
    def _detect_file_type_by_content(self, file_path: str) -> str:
        """Определить тип файла по содержимому, а не только расширению"""
        try:
            mime_type = magic.from_file(file_path, mime=True)
            logger.info(f"Detected MIME type for {file_path}: {mime_type}")
            
            # Маппинг MIME-типов к расширениям
            mime_to_ext = {
                'application/pdf': '.pdf',
                'application/msword': '.doc',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                'application/vnd.ms-excel': '.xls',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
                'text/plain': '.txt',
                'image/jpeg': '.jpg',
                'image/png': '.png',
                'image/tiff': '.tiff',
                'image/bmp': '.bmp'
            }
            detected_ext = mime_to_ext.get(mime_type, os.path.splitext(file_path)[1].lower())
            logger.info(f"Mapped to extension: {detected_ext}")
            return detected_ext
            
        except Exception as e:
            logger.warning(f"Failed to detect MIME type for {file_path}: {e}")
            return os.path.splitext(file_path)[1].lower()
    
    def _create_user_friendly_error(self, error: Exception, filename: str) -> dict:
        """Создать понятное сообщение об ошибке для пользователя"""
        error_str = str(error)
        
        # Определить тип ошибки и дать понятное объяснение
        if "not a Word file" in error_str or "not a valid" in error_str.lower():
            user_message = f"Файл '{filename}' поврежден или имеет неправильное расширение. Попросите поставщика пересохранить файл в правильном формате."
            suggestion = "Рекомендуем использовать форматы: PDF, DOCX, XLSX"
        elif "corrupted" in error_str.lower() or "damaged" in error_str.lower():
            user_message = f"Файл '{filename}' поврежден и не может быть прочитан"
            suggestion = "Попросите поставщика отправить файл заново"
        elif "password" in error_str.lower() or "encrypted" in error_str.lower():
            user_message = f"Файл '{filename}' защищен паролем"
            suggestion = "Попросите поставщика убрать защиту или отправить пароль"
        elif "permission" in error_str.lower() or "access" in error_str.lower():
            user_message = f"Нет доступа к файлу '{filename}'"
            suggestion = "Проверьте права доступа к файлу"
        else:
            user_message = f"Не удалось обработать файл '{filename}'"
            suggestion = "Попробуйте другой формат файла (PDF, DOCX, XLSX)"
        
        return {
            "user_message": user_message,
            "suggestion": suggestion,
            "technical_error": error_str,
            "status": "partial_failure"
        }
    
    def _process_with_fallback(self, temp_file_path: str, filename: str) -> tuple:
        """Попробовать несколько методов обработки и вернуть результат + статус"""
        logger.info(f"Starting fallback processing for {filename}")
        
        # 1. Определить тип по содержимому
        detected_type = self._detect_file_type_by_content(temp_file_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        logger.info(f"File extension: {file_ext}, Detected type: {detected_type}")
        
        # 2. Список методов для попытки
        methods_to_try = []
        
        # Основной метод на основе определенного типа
        if detected_type in ['.pdf']:
            methods_to_try.append(('PDF (content-based)', self._process_pdf))
        elif detected_type in ['.doc', '.docx']:
            # For .doc files, try multiple methods in order of preference
            if file_ext == '.doc':
                methods_to_try.append(('Old Word (.doc)', self._process_old_word))
                methods_to_try.append(('LibreOffice DOC→DOCX', self._process_doc_with_libreoffice))
            methods_to_try.append(('Word (content-based)', self._process_word))
        elif detected_type in ['.xls', '.xlsx']:
            methods_to_try.append(('Excel (content-based)', self._process_excel))
        elif detected_type in ['.txt']:
            methods_to_try.append(('Text (content-based)', self._process_text))
        
        # Fallback методы если основной тип не совпадает с расширением
        if file_ext != detected_type:
            logger.info(f"Extension mismatch detected. Adding fallback methods.")
            
            if file_ext in ['.doc', '.docx'] and detected_type not in ['.doc', '.docx']:
                # For .doc files, try all available methods
                if file_ext == '.doc':
                    methods_to_try.append(('Old Word (.doc extension-based)', self._process_old_word))
                    methods_to_try.append(('LibreOffice DOC→DOCX (extension-based)', self._process_doc_with_libreoffice))
                methods_to_try.append(('Word (extension-based)', self._process_word))
            if file_ext in ['.pdf'] and detected_type != '.pdf':
                methods_to_try.append(('PDF (extension-based)', self._process_pdf))
            if file_ext in ['.xls', '.xlsx'] and detected_type not in ['.xls', '.xlsx']:
                methods_to_try.append(('Excel (extension-based)', self._process_excel))
            if file_ext in ['.txt'] and detected_type != '.txt':
                methods_to_try.append(('Text (extension-based)', self._process_text))
        
        # Text reading as fallback before OCR (for corrupted docs that are actually text/XML)
        if detected_type in ['.txt'] or 'text' in detected_type or 'xml' in detected_type:
            methods_to_try.append(('Text reading (corrupted doc)', self._process_text))
        
        # OCR как последний fallback для всех файлов (only for actual images)
        methods_to_try.append(('OCR (last resort)', self._process_image))
        
        logger.info(f"Will try {len(methods_to_try)} methods: {[m[0].encode('ascii', 'replace').decode('ascii') for m in methods_to_try]}")
        
        # 3. Попробовать методы по очереди
        last_error = None
        for method_name, method_func in methods_to_try:
            try:
                logger.info(f"Trying method: {method_name}")
                result = method_func(temp_file_path)
                
                # Проверить что результат содержательный
                if result and str(result).strip():
                    # Для Excel результат - это dict, для остальных - string
                    if isinstance(result, dict):
                        text_result = self._excel_data_to_text(result)
                        success = len(text_result.strip()) > 0
                        final_result = result
                    else:
                        success = len(result.strip()) > 0
                        text_result = result
                        final_result = result
                    
                    if success:
                        logger.info(f"Successfully processed {filename} using {method_name}")
                        return final_result, {"status": "success", "method": method_name}
                        
                logger.warning(f"Method {method_name} returned empty result for {filename}")
                
            except Exception as e:
                logger.warning(f"Method {method_name} failed for {filename}: {e}")
                last_error = e
                continue
        
        # Все методы не сработали
        logger.error(f"All processing methods failed for {filename}")
        error_info = self._create_user_friendly_error(
            last_error or Exception("All processing methods failed"), filename
        )
        return f"Ошибка обработки: {error_info['user_message']}", error_info
    
    def _process_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        logger.info(f"Processing text file: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'cp1251', 'latin1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"Successfully read text file with {encoding} encoding: {len(content)} chars")
                    return content
                except UnicodeDecodeError:
                    continue
                    
            # If all encodings fail, try binary mode
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Try to decode as UTF-8 with errors handling
            text_content = content.decode('utf-8', errors='replace')
            logger.info(f"Read text file in binary mode with error replacement: {len(text_content)} chars")
            return text_content
            
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise
    
    def _process_image(self, file_path: str) -> str:
        """
        Extract text from image using advanced OCR with multiple preprocessing techniques
        to handle various types of document images.
        """
        logger.info(f"Processing image with OCR: {file_path}")
        
        # First, check if this might be a text file disguised as something else
        try:
            # Try to read as text first - if it's actually text content, don't use OCR
            with open(file_path, 'rb') as f:
                content = f.read(1024)  # Read first 1KB
            
            # Check if content looks like text (XML, HTML, plain text, etc.)
            try:
                text_content = content.decode('utf-8', errors='ignore')
                if len(text_content) > 10 and any(c.isalnum() for c in text_content):
                    # If it contains readable text, try text processing instead
                    if any(marker in text_content.lower() for marker in ['<?xml', '<html', 'text', 'document']):
                        logger.info(f"File appears to be text-based, trying text extraction instead of OCR")
                        return self._process_text(file_path)
            except:
                pass
        except:
            pass
        print(f"[IMAGE] Processing image: {file_path}")
        
        try:
            # Create a temporary directory for preprocessed images
            temp_dir = tempfile.mkdtemp()
            print(f"[IMAGE] Created temporary directory: {temp_dir}")
            
            # Load the image with OpenCV
            img = cv2.imread(file_path)
            
            # Check if image was loaded properly
            if img is None:
                logger.warning(f"Could not load image with OpenCV: {file_path}")
                print(f"[IMAGE] WARNING: Could not load image with OpenCV, trying with PIL")
                
                # Try with PIL as a fallback
                pil_img = Image.open(file_path)
                text = pytesseract.image_to_string(pil_img, config=self.tesseract_config)
                print(f"[IMAGE] Extracted {len(text)} chars with PIL fallback")
                return text
            
            # Try multiple preprocessing techniques and use the best result
            print(f"[IMAGE] Applying multiple preprocessing techniques")
            preprocessing_results = []
            
            # 1. Original image (baseline)
            print(f"[IMAGE] Method 1: Original image")
            original_text = pytesseract.image_to_string(Image.open(file_path), config=self.tesseract_config)
            preprocessing_results.append((original_text, "original"))
            
            # 2. Basic grayscale + threshold
            print(f"[IMAGE] Method 2: Basic grayscale + threshold")
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            # Apply threshold
            _, threshold = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            threshold_path = os.path.join(temp_dir, "threshold.png")
            cv2.imwrite(threshold_path, threshold)
            threshold_text = pytesseract.image_to_string(Image.open(threshold_path), config=self.tesseract_config)
            preprocessing_results.append((threshold_text, "threshold"))
            
            # 3. Enhanced contrast with CLAHE
            print(f"[IMAGE] Method 3: Enhanced contrast with CLAHE")
            # Apply adaptive equalization to enhance contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            contrast_enhanced = clahe.apply(gray)
            # Apply threshold
            _, enhanced_threshold = cv2.threshold(contrast_enhanced, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            enhanced_path = os.path.join(temp_dir, "enhanced.png")
            cv2.imwrite(enhanced_path, enhanced_threshold)
            enhanced_text = pytesseract.image_to_string(Image.open(enhanced_path), config=self.tesseract_config)
            preprocessing_results.append((enhanced_text, "enhanced"))
            
            # 4. Denoising
            print(f"[IMAGE] Method 4: Denoising")
            # Apply fastNlMeansDenoising
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            _, denoised_threshold = cv2.threshold(denoised, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            denoised_path = os.path.join(temp_dir, "denoised.png")
            cv2.imwrite(denoised_path, denoised_threshold)
            denoised_text = pytesseract.image_to_string(Image.open(denoised_path), config=self.tesseract_config)
            preprocessing_results.append((denoised_text, "denoised"))
            
            # 5. Sharpening with unsharp mask
            print(f"[IMAGE] Method 5: Sharpening")
            # Create a gaussian blur and apply unsharp mask
            gaussian = cv2.GaussianBlur(gray, (0, 0), 3)
            sharpened = cv2.addWeighted(gray, 1.5, gaussian, -0.5, 0)
            _, sharpened_threshold = cv2.threshold(sharpened, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            sharpened_path = os.path.join(temp_dir, "sharpened.png")
            cv2.imwrite(sharpened_path, sharpened_threshold)
            sharpened_text = pytesseract.image_to_string(Image.open(sharpened_path), config=self.tesseract_config)
            preprocessing_results.append((sharpened_text, "sharpened"))
            
            # 6. Color separation (for colored backgrounds)
            print(f"[IMAGE] Method 6: Color channel separation")
            b, g, r = cv2.split(img)
            # Process all channels
            channels = {"blue": b, "green": g, "red": r}
            for name, channel in channels.items():
                _, channel_threshold = cv2.threshold(channel, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                channel_path = os.path.join(temp_dir, f"{name}_channel.png")
                cv2.imwrite(channel_path, channel_threshold)
                channel_text = pytesseract.image_to_string(Image.open(channel_path), config=self.tesseract_config)
                preprocessing_results.append((channel_text, f"{name}_channel"))
            
            # 7. Scale image up (can improve OCR for small text)
            print(f"[IMAGE] Method 7: Image scaling")
            h, w = gray.shape
            scaled = cv2.resize(gray, (w*2, h*2), interpolation=cv2.INTER_CUBIC)
            _, scaled_threshold = cv2.threshold(scaled, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            scaled_path = os.path.join(temp_dir, "scaled.png")
            cv2.imwrite(scaled_path, scaled_threshold)
            scaled_text = pytesseract.image_to_string(Image.open(scaled_path), config=self.tesseract_config)
            preprocessing_results.append((scaled_text, "scaled"))
            
            # Compare results and choose the best one (most text content)
            valid_results = [(text, method) for text, method in preprocessing_results if text and text.strip()]
            
            if valid_results:
                # Sort by text length and choose the longest result
                valid_results.sort(key=lambda x: len(x[0]), reverse=True)
                best_text, best_method = valid_results[0]
                
                # Log stats about the extraction
                text_length = len(best_text)
                logger.info(f"Successfully extracted {text_length} chars with method '{best_method}'")
                print(f"[IMAGE] Best result from method '{best_method}' with {text_length} chars")
                
                # Print sample of extracted text
                sample_length = min(100, text_length)
                print(f"[IMAGE] Sample text: {best_text[:sample_length]}...")
                
                # Clean up text
                cleaned_text = re.sub(r'\n{3,}', '\n\n', best_text)  # Replace 3+ newlines with 2
                cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)   # Replace multiple spaces with one
                
                return cleaned_text
            else:
                logger.warning("All preprocessing methods failed to extract text")
                print(f"[IMAGE] All preprocessing methods failed to extract text")
                return ""
                
        except Exception as e:
            logger.error(f"Error in image processing: {str(e)}")
            print(f"[IMAGE] ERROR: Image processing failed: {str(e)}")
            traceback_str = traceback.format_exc()
            logger.error(traceback_str)
            print(f"[IMAGE] Traceback: {traceback_str}")
            raise
            
        finally:
            # Clean up temp directory if it exists
            if 'temp_dir' in locals() and temp_dir and os.path.exists(temp_dir):
                print(f"[IMAGE] Cleaning up temporary directory: {temp_dir}")
                shutil.rmtree(temp_dir, ignore_errors=True)
    
    def _save_processing_result(self, result: Dict[str, Any], filename: str, email: str) -> None:
        """Save processing result to disk"""
        # Sanitize filename to make it safe for filesystem
        safe_filename = ''.join(c if c.isalnum() or c in '._- ' else '_' for c in filename)
        safe_email = ''.join(c if c.isalnum() or c in '._-@' else '_' for c in email)
        
        # Create a unique output filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        output_filename = f"{safe_email}_{safe_filename}_{timestamp}.json"
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Save as JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2, default=str)
        
        logger.info(f"Saved processing result to: {output_path}")


# Function to batch process attachments from the database
def process_attachment_batch(attachments: List[Dict[str, Any]], output_dir: str = "processed_attachments") -> Dict[str, Any]:
    """
    Process a batch of attachments
    
    Args:
        attachments: List of attachment dictionaries with:
            - content: Binary file content
            - filename: Name of the file
            - supplier_name: Name of the supplier
            - supplier_email: Email address of the supplier
            - email_date: Date of the email
            - request_id: Optional ID of the associated request
        output_dir: Directory to store processed output
        
    Returns:
        Dictionary with processing statistics
    """
    processor = FileProcessor(output_dir=output_dir)
    
    stats = {
        "total": len(attachments),
        "successful": 0,
        "failed": 0,
        "by_type": {}
    }
    
    results = []
    
    for attachment in attachments:
        try:
            result = processor.process_attachment(
                file_data=attachment.get('content', b''),
                filename=attachment.get('filename', 'unknown.file'),
                supplier_name=attachment.get('supplier_name', 'Unknown Supplier'),
                supplier_email=attachment.get('supplier_email', 'unknown@example.com'),
                email_date=attachment.get('email_date', datetime.now().isoformat()),
                request_id=attachment.get('request_id')
            )
            
            # Update statistics
            stats["successful"] += 1
            
            # Track file types
            file_ext = os.path.splitext(attachment.get('filename', ''))[1].lower().lstrip('.')
            if not file_ext:
                file_ext = 'unknown'
                
            if file_ext not in stats["by_type"]:
                stats["by_type"][file_ext] = {"count": 0, "success": 0, "failed": 0}
            
            stats["by_type"][file_ext]["count"] += 1
            
            if result.get("metadata", {}).get("error", False):
                stats["by_type"][file_ext]["failed"] += 1
            else:
                stats["by_type"][file_ext]["success"] += 1
                
            results.append(result)
            
        except Exception as e:
            logger.error(f"Error processing attachment {attachment.get('filename', 'unknown')}: {str(e)}")
            stats["failed"] += 1
            
            # Track file types for failures too
            file_ext = os.path.splitext(attachment.get('filename', ''))[1].lower().lstrip('.')
            if not file_ext:
                file_ext = 'unknown'
                
            if file_ext not in stats["by_type"]:
                stats["by_type"][file_ext] = {"count": 0, "success": 0, "failed": 0}
            
            stats["by_type"][file_ext]["count"] += 1
            stats["by_type"][file_ext]["failed"] += 1
    
    # Save summary report
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    summary_path = os.path.join(output_dir, f"summary_{timestamp}.json")
    with open(summary_path, 'w', encoding='utf-8') as f:
        summary = {
            "stats": stats,
            "timestamp": datetime.now().isoformat(),
            "results": results
        }
        json.dump(summary, f, ensure_ascii=False, indent=2, default=str)
    
    logger.info(f"Batch processing complete: {stats['successful']} successful, {stats['failed']} failed")
    return {"stats": stats, "results": results}


if __name__ == "__main__":
    # This can be used for command-line testing
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        processor = FileProcessor()
        
        with open(file_path, 'rb') as f:
            file_data = f.read()
            
        result = processor.process_attachment(
            file_data=file_data,
            filename=os.path.basename(file_path),
            supplier_name="Test Supplier",
            supplier_email="test@example.com",
            email_date=datetime.now().isoformat()
        )
        
        print(json.dumps(result, indent=2, default=str))