#!/usr/bin/env python3
"""
Simple Text Extractor for SupplierFinder
Based on the working extract_text_optimized.py from checkpoint 142ef8e75672e4808d19f33c8adbb292c7442f4a
But adapted for backend processing with JSON output
"""

import sys
import os
import json
import tempfile
import subprocess
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    # Tesseract должен быть в PATH - не устанавливаем принудительный путь
    from docx import Document
    import pandas as pd
    import openpyxl
    import magic
    import chardet
    import easyocr
    import striprtf
except ImportError as e:
    print(json.dumps({"error": f"Missing required package: {e}"}, ensure_ascii=False))
    sys.exit(1)

def extract_text_from_pdf(file_path):
    """
    Извлекает текст из PDF файлов
    Сначала пытается извлечь текст через PyPDF2, затем использует OCR для отсканированных PDF
    """
    text = ""
    
    try:
        # Шаг 1: Пытаемся извлечь текст через PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n[PAGE:{page_num}]\n{page_text}\n"
                except Exception as e:
                    continue
        
        # Шаг 2: Если текст не извлечен, используем OCR
        if not text.strip():
            try:
                # Конвертируем PDF в изображения
                images = convert_from_path(file_path)
                
                # Обрабатываем каждую страницу с помощью OCR
                for i, image in enumerate(images, 1):
                    page_text = ""
                    
                    # Пробуем разные PSM режимы и языки для лучшего распознавания
                    configs = [
                        ('rus+eng', '--oem 3 --psm 1'),
                        ('rus+eng', '--oem 3 --psm 3'),
                        ('rus+eng', '--oem 3 --psm 6'),
                        ('rus+eng', '--oem 3 --psm 8'),
                        ('rus', '--oem 3 --psm 1'),
                        ('eng', '--oem 3 --psm 1'),
                    ]
                    
                    # Сначала пробуем EasyOCR для страницы
                    temp_path = f"temp_page_{i}.png"
                    try:
                        image.save(temp_path, 'PNG')  # Явно указываем формат PNG
                        
                        easyocr_text = extract_text_with_easyocr(temp_path)
                        if easyocr_text and not easyocr_text.startswith("Error") and len(easyocr_text.strip()) > 10:
                            page_text = easyocr_text
                        else:
                            # Fallback на Tesseract с разными конфигурациями
                            for lang, config in configs:
                                try:
                                    # Предобработка изображения для лучшего OCR
                                    processed_image = preprocess_image_for_ocr(image)
                                    
                                    # Извлекаем текст с помощью pytesseract
                                    page_text = pytesseract.image_to_string(
                                        processed_image, 
                                        lang=lang,
                                        config=config
                                    )
                                    
                                    # Убеждаемся, что текст в правильной кодировке
                                    if isinstance(page_text, bytes):
                                        page_text = page_text.decode('utf-8', errors='replace')
                                    
                                    # Если получили осмысленный текст, используем его
                                    if page_text.strip() and len(page_text.strip()) > 10:
                                        # Проверяем, что в тексте есть русские символы или цифры
                                        has_meaningful_content = any(
                                            char.isalpha() or char.isdigit() or char in '.,!?;:()[]{}"\'«»№%+-= '
                                            for char in page_text
                                        )
                                        if has_meaningful_content:
                                            break
                                        
                                except Exception as config_error:
                                    continue
                    except Exception as temp_error:
                        # Если не удалось сохранить временный файл, используем только Tesseract
                        page_text = ""
                        for lang, config in configs:
                            try:
                                # Предобработка изображения для лучшего OCR
                                processed_image = preprocess_image_for_ocr(image)
                                
                                # Извлекаем текст с помощью pytesseract
                                page_text = pytesseract.image_to_string(
                                    processed_image, 
                                    lang=lang,
                                    config=config
                                )
                                
                                # Убеждаемся, что текст в правильной кодировке
                                if isinstance(page_text, bytes):
                                    page_text = page_text.decode('utf-8', errors='replace')
                                
                                # Если получили осмысленный текст, используем его
                                if page_text.strip() and len(page_text.strip()) > 10:
                                    # Проверяем, что в тексте есть русские символы или цифры
                                    has_meaningful_content = any(
                                        char.isalpha() or char.isdigit() or char in '.,!?;:()[]{}"\'«»№%+-= '
                                        for char in page_text
                                    )
                                    if has_meaningful_content:
                                        break
                                    
                            except Exception as config_error:
                                continue
                    finally:
                        # Удаляем временный файл
                        if os.path.exists(temp_path):
                            try:
                                os.remove(temp_path)
                            except Exception:
                                pass  # Игнорируем ошибки удаления
                    
                    if page_text.strip():
                        text += f"\n[PAGE:{i}]\n{page_text}\n"
                        
            except Exception as ocr_error:
                # Если OCR не удался, возвращаем то, что удалось извлечь
                pass
    
    except Exception as e:
        return f"Error extracting from PDF: {e}"
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX"""
    text = ""
    
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text += para_text + "\n"
        
        # Extract tables
        for table in doc.tables:
            text += "\n[TABLE]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "[/TABLE]\n"
    
    except Exception as e:
        return f"Error extracting from DOCX: {e}"
    
    return text

def extract_text_from_txt(file_path):
    """Extract text from plain text files with encoding detection"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'latin1', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use binary mode with error replacement
        with open(file_path, 'rb') as file:
            content = file.read()
            return content.decode('utf-8', errors='replace')
    
    except Exception as e:
        return f"Error extracting from text file: {e}"

def extract_text_from_csv(file_path):
    """Extract text from CSV files using pandas with multiple encoding support"""
    try:
        # First try UTF-8 encoding
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
            return df.to_string()
        except UnicodeDecodeError:
            # If UTF-8 fails, try cp1251 (common for Windows)
            try:
                df = pd.read_csv(file_path, encoding='cp1251')
                return df.to_string()
            except UnicodeDecodeError:
                # If cp1251 fails, try latin-1
                try:
                    df = pd.read_csv(file_path, encoding='latin-1')
                    return df.to_string()
                except Exception as e:
                    return f"Error reading CSV with any encoding: {e}"
    except Exception as e:
        return f"Error extracting from CSV: {e}"

def extract_text_from_excel(file_path):
    """Extract text from Excel files"""
    try:
        # Determine engine based on file extension
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.xls':
            engine = 'xlrd'
        else:
            engine = 'openpyxl'
        
        # Read all sheets
        excel_data = {}
        xls = pd.ExcelFile(file_path, engine=engine)
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            
            # Convert to text format
            text = f"\n[SHEET: {sheet_name}]\n"
            
            # Add headers
            if not df.columns.empty:
                text += " | ".join(str(h) for h in df.columns) + "\n"
                text += "-" * 50 + "\n"
            
            # Add data rows (limit to 100 rows)
            for i, row in df.iterrows():
                if i >= 100:
                    text += f"... and {len(df) - 100} more rows\n"
                    break
                text += " | ".join(str(cell) for cell in row) + "\n"
            
            excel_data[sheet_name] = text
        
        return "\n".join(excel_data.values())
    
    except Exception as e:
        return f"Error extracting from Excel: {e}"


def preprocess_image_for_ocr(image):
    """
    Предобработка изображения для улучшения качества OCR
    """
    try:
        import cv2
        import numpy as np
        
        # Конвертируем PIL Image в numpy array для OpenCV
        if hasattr(image, 'mode'):
            # PIL Image
            img_array = np.array(image)
            if len(img_array.shape) == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        else:
            # Уже numpy array
            img_array = image
        
        # Преобразование в оттенки серого
        gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
        
        # Увеличение контраста
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        
        # Удаление шума
        denoised = cv2.medianBlur(enhanced, 3)
        
        # Морфологические операции для улучшения текста
        kernel = np.ones((1,1), np.uint8)
        processed = cv2.morphologyEx(denoised, cv2.MORPH_CLOSE, kernel)
        
        # Адаптивная бинаризация
        binary = cv2.adaptiveThreshold(processed, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # Конвертируем обратно в PIL Image
        processed_image = Image.fromarray(binary)
        
        return processed_image
        
    except Exception as e:
        # Если предобработка не удалась, возвращаем оригинальное изображение
        return image

def extract_text_with_easyocr(file_path):
    """
    Извлекает текст с помощью EasyOCR (лучше для русского текста)
    """
    try:
        import easyocr
        # Инициализируем EasyOCR с поддержкой русского и английского
        reader = easyocr.Reader(['ru', 'en'], gpu=False)
        
        # Читаем изображение
        results = reader.readtext(file_path)
        
        # Объединяем все найденные тексты
        text_parts = []
        for (bbox, text, confidence) in results:
            if confidence > 0.3:  # Более низкий порог для лучшего распознавания
                text_parts.append(text)
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        return f"Error with EasyOCR: {e}"

def extract_text_from_image(file_path):
    """
    Извлекает текст из изображений (JPG, PNG) с использованием EasyOCR и Tesseract
    Поддерживает русский и английский языки
    """
    try:
        # Сначала пробуем EasyOCR (лучше для русского текста)
        easyocr_text = extract_text_with_easyocr(file_path)
        
        # Если EasyOCR дал хороший результат, используем его
        if easyocr_text and len(easyocr_text.strip()) > 5 and not easyocr_text.startswith("Error"):
            return easyocr_text
        
        # Иначе пробуем Tesseract как fallback
        image = Image.open(file_path)
        
        # Пробуем разные конфигурации Tesseract
        configs = [
            ('rus+eng', '--oem 3 --psm 6'),
            ('rus+eng', '--oem 3 --psm 1'),
            ('rus+eng', '--oem 3 --psm 3'),
            ('rus', '--oem 3 --psm 6'),
            ('eng', '--oem 3 --psm 6'),
        ]
        
        best_tesseract_text = ""
        for lang, config in configs:
            try:
                processed_image = preprocess_image_for_ocr(image)
                
                tesseract_text = pytesseract.image_to_string(
                    processed_image, 
                    lang=lang,
                    config=config
                )
                
                # Убеждаемся, что текст в правильной кодировке
                if isinstance(tesseract_text, bytes):
                    tesseract_text = tesseract_text.decode('utf-8', errors='replace')
                
                # Если получили осмысленный текст, используем его
                if tesseract_text.strip() and len(tesseract_text.strip()) > len(best_tesseract_text.strip()):
                    # Проверяем, что в тексте есть русские символы или цифры
                    has_meaningful_content = any(
                        char.isalpha() or char.isdigit() or char in '.,!?;:()[]{}"\'«»№%+-= '
                        for char in tesseract_text
                    )
                    if has_meaningful_content:
                        best_tesseract_text = tesseract_text.strip()
                        
            except Exception as config_error:
                continue
        
        # Выбираем лучший результат
        if len(best_tesseract_text.strip()) > len(easyocr_text.strip()):
            return best_tesseract_text
        
        return easyocr_text if easyocr_text else best_tesseract_text
    
    except Exception as e:
        return f"Error extracting from image: {e}"

def extract_text_from_rtf(rtf_path):
    """
    Извлекает текст из RTF файла.
    """
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as file:
            rtf_content = file.read()
            from striprtf.striprtf import rtf_to_text
            text = rtf_to_text(rtf_content)
            return text.strip()
    except Exception as e:
        return f"ERROR: Не удалось обработать RTF файл {rtf_path}. Ошибка: {e}"

def find_soffice_path():
    """Динамически ищет soffice.exe в стандартных директориях Windows."""
    possible_paths = [
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe"
    ]
    for path in possible_paths:
        if os.path.exists(path):
            print(f"DEBUG: soffice.exe найден по пути: {path}", file=sys.stderr)
            return path
    print("DEBUG: soffice.exe не найден ни по одному из стандартных путей.", file=sys.stderr)
    return None

def convert_with_libreoffice(doc_path, output_dir):
    """Конвертирует документ, используя динамически найденный путь к soffice."""
    print("DEBUG: Начало конвертации LibreOffice", file=sys.stderr)
    
    # --- ШАГ 1: ПОИСК SOFFICE.EXE ---
    print("DEBUG: [ШАГ 1] Поиск исполняемого файла soffice.exe...", file=sys.stderr)
    soffice_executable_path = find_soffice_path()
    
    if not soffice_executable_path:
        print("ERROR: [ШАГ 1] LibreOffice (soffice) не найден. Этот метод будет пропущен.", file=sys.stderr)
        return None

    # --- ШАГ 2: ПОДГОТОВКА ФАЙЛОВ ---
    print("DEBUG: [ШАГ 2] Подготовка файлов...", file=sys.stderr)
    print(f"DEBUG: Входной файл: {doc_path}", file=sys.stderr)
    print(f"DEBUG: Выходная директория: {output_dir}", file=sys.stderr)
    
    if not os.path.exists(doc_path):
        print(f"ERROR: [ШАГ 2] Входной файл не найден: {doc_path}", file=sys.stderr)
        return None

    # --- ШАГ 3: ВЫПОЛНЕНИЕ КОМАНДЫ ---
    print("DEBUG: [ШАГ 3] Выполнение команды конвертации...", file=sys.stderr)
    command = [
        soffice_executable_path,
        "--headless",
        "--convert-to", "docx",
        "--outdir", output_dir,
        doc_path
    ]
    print(f"DEBUG: Выполняемая команда: {' '.join(command)}", file=sys.stderr)

    try:
        result = subprocess.run(
            command, 
            check=True,        # Обязательно, чтобы выбросить исключение при ошибке
            timeout=60, 
            capture_output=True, # Захватываем вывод
            text=True            # Декодируем в текст
        )
        print("DEBUG: [ШАГ 3] УСПЕХ! Команда выполнилась без ошибок.", file=sys.stderr)
        print(f"DEBUG: Вывод LibreOffice (stdout): {result.stdout or 'пусто'}", file=sys.stderr)
        print(f"DEBUG: Ошибки LibreOffice (stderr): {result.stderr or 'пусто'}", file=sys.stderr)

    except FileNotFoundError:
        print("ERROR: [ШАГ 3] КРИТИЧЕСКАЯ ОШИБКА: Система не нашла команду, хотя Python видит файл.", file=sys.stderr)
        return None
    except subprocess.CalledProcessError as e:
        print(f"ERROR: [ШАГ 3] КРИТИЧЕСКАЯ ОШИБКА: LibreOffice вернул код ошибки {e.returncode}.", file=sys.stderr)
        print(f"ERROR: Вывод (stdout): {e.stdout or 'пусто'}", file=sys.stderr)
        print(f"ERROR: Ошибки (stderr): {e.stderr or 'пусто'}", file=sys.stderr)
        return None
    except subprocess.TimeoutExpired:
        print("ERROR: [ШАГ 3] КРИТИЧЕСКАЯ ОШИБКА: LibreOffice превысил таймаут (60 сек).", file=sys.stderr)
        return None
    except Exception as e:
        print(f"ERROR: [ШАГ 3] КРИТИЧЕСКАЯ ОШИБКА: Произошла непредвиденная ошибка: {e}", file=sys.stderr)
        return None

    # --- ШАГ 4: ПРОВЕРКА РЕЗУЛЬТАТА ---
    print("DEBUG: [ШАГ 4] Проверка созданного .docx файла...", file=sys.stderr)
    base_name = os.path.basename(doc_path)
    new_name = os.path.splitext(base_name)[0] + ".docx"
    new_path = os.path.join(output_dir, new_name)
    
    if os.path.exists(new_path):
        file_size = os.path.getsize(new_path)
        print(f"INFO: [ШАГ 4] УСПЕХ! Сконвертированный файл найден: {new_path}", file=sys.stderr)
        print(f"INFO: Размер файла: {file_size} байт.", file=sys.stderr)
        return new_path
    else:
        print("ERROR: [ШАГ 4] КРИТИЧЕСКАЯ ОШИБКА: Команда отработала, но итоговый файл не был создан.", file=sys.stderr)
        return None

def process_complex_doc(doc_path, temp_dir):
    """Главный обработчик, использующий LibreOffice."""
    print("DEBUG: Вход в process_complex_doc.", file=sys.stderr)
    
    converted_path = convert_with_libreoffice(doc_path, temp_dir)
    
    if converted_path and os.path.exists(converted_path):
        text = extract_text_from_docx(converted_path)
        os.remove(converted_path)
        if text:
            print("INFO: Текст успешно извлечен после конвертации LibreOffice.", file=sys.stderr)
            return text
    
    print("INFO: Метод LibreOffice не сработал. Переход к резервным методам...", file=sys.stderr)
    # Используем существующий fallback метод
    return extract_text_from_doc_improved(doc_path)

def extract_text_from_doc_improved(doc_path):
    """
    Извлекает текст из .doc файлов.
    Основной метод: 'antiword' для максимальной совместимости.
    Резервный метод (Fallback): 'docx2txt' на случай, если antiword не установлен.
    """
    if not os.path.exists(doc_path):
        return f"ERROR: Файл не найден: {doc_path}"

    # --- Метод 1: Попытка использовать 'antiword' (предпочтительный) ---
    try:
        # Проверяем наличие antiword без вывода ошибок в консоль
        subprocess.run(['antiword', '-h'], check=True, capture_output=True)
        
        # Если antiword найден, используем его
        process = subprocess.run(
            ['antiword', doc_path],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore'
        )
        if process.returncode == 0 and process.stdout.strip():
            return process.stdout.strip()
        else:
            pass  # Переходим к fallback

    except (FileNotFoundError, subprocess.CalledProcessError):
        # Эта ошибка означает, что antiword не найден или не работает
        pass

    # --- Метод 2: Fallback с использованием 'docx2txt' ---
    try:
        import docx2txt
        text = docx2txt.process(doc_path)
        if text and text.strip():
            return text.strip()
        else:
            return f"ERROR: Fallback-метод (docx2txt) не смог извлечь текст из {doc_path}."
    except Exception as e:
        return f"ERROR: Не удалось обработать {doc_path}. Основной метод (antiword) и fallback (docx2txt) не сработали. Ошибка: {e}"

def extract_text_from_old_doc(file_path):
    """Extract text from old DOC files using multiple methods"""
    try:
        # Method 1: Try docx2txt (can handle some .doc files)
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and len(text.strip()) > 10:
                return text.strip()
        except Exception:
            pass
        
        # Method 2: Try reading as binary and extracting text manually
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Try different encodings
            encodings = ['utf-8', 'windows-1251', 'cp1251', 'latin1']
            text_parts = []
            
            for encoding in encodings:
                try:
                    decoded = content.decode(encoding, errors='ignore')
                    # Extract readable text
                    import re
                    readable_text = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()\-]+', decoded)
                    if readable_text:
                        text_parts.extend(readable_text)
                except:
                    continue
            
            if text_parts:
                combined_text = ' '.join(text_parts)
                # Clean up
                import re
                cleaned_text = re.sub(r'\s+', ' ', combined_text).strip()
                return cleaned_text
            
            return ""
            
        except Exception:
            pass
        
        # Method 3: Try LibreOffice conversion (if available)
        try:
            import subprocess
            import tempfile
            import shutil
            
            temp_dir = tempfile.mkdtemp()
            try:
                # Use LibreOffice to convert DOC to DOCX
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, file_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # Find the converted DOCX file
                    base_name = Path(file_path).stem
                    converted_file = Path(temp_dir) / f"{base_name}.docx"
                    
                    if converted_file.exists():
                        # Process the converted DOCX file
                        doc = Document(converted_file)
                        full_text = []
                        
                        for para in doc.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text.strip())
                        
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    if cell.text.strip():
                                        row_text.append(cell.text.strip())
                                if row_text:
                                    full_text.append(" | ".join(row_text))
                        
                        return "\n".join(full_text)
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
        
        return "Не удалось извлечь текст из DOC файла"
    
    except Exception as e:
        return f"Error extracting from DOC: {e}"

def detect_file_type(file_path):
    """Detect file type using python-magic with improved fallback"""
    try:
        # Use python-magic to detect MIME type from file content
        mime_type = magic.from_file(file_path, mime=True)
        if mime_type and mime_type != 'application/octet-stream':
            return mime_type
    except Exception as e:
        # If magic fails, log the error but continue
        pass
    
    # Fallback to file extension
    file_ext = Path(file_path).suffix.lower()
    ext_to_mime = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.rtf': 'application/rtf',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.txt': 'text/plain',
        '.csv': 'text/csv',
        '.html': 'text/html',
        '.xml': 'text/xml'
    }
    return ext_to_mime.get(file_ext, 'application/octet-stream')

def process_file_smart(file_path):
    """
    Определяет реальный тип файла по его содержимому (MIME-type) 
    и вызывает соответствующий обработчик.
    """
    if not os.path.exists(file_path):
        return "ERROR: Файл не найден."

    try:
        # Определяем MIME-тип файла
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        print(f"INFO: Обнаружен тип файла '{file_type}' для {os.path.basename(file_path)}", file=sys.stderr)

    except Exception as e:
        return f"ERROR: Не удалось определить тип файла с помощью python-magic: {e}"
    
    # --- Маршрутизация на основе реального типа файла ---

    if 'text/rtf' in file_type:
        return extract_text_from_rtf(file_path)
    
    elif 'pdf' in file_type:
        return extract_text_from_pdf(file_path) # Функция для PDF
    
    elif 'msword' in file_type: # Для .doc и некоторых .docx
        # Используем новый обработчик с LibreOffice приоритетом
        return process_complex_doc(file_path, tempfile.gettempdir()) 
    
    elif 'openxmlformats-officedocument.wordprocessingml' in file_type: # Для .docx
        return extract_text_from_docx(file_path)
    
    elif 'vnd.ms-excel' in file_type or 'openxmlformats-officedocument.spreadsheetml' in file_type:
        return extract_text_from_excel(file_path) # Функция для Excel
        
    elif 'text' in file_type: # Для .txt, .csv, .html, .xml
        # Для CSV лучше оставить отдельную логику, если она нужна
        if file_path.endswith('.csv'):
             return extract_text_from_csv(file_path)
        return extract_text_from_txt(file_path)
        
    elif 'image' in file_type:
        return extract_text_from_image(file_path) # OCR функция для картинок

    else:
        # Если тип не распознан, пробуем по расширению как раньше
        print(f"WARNING: Неизвестный MIME-тип '{file_type}'. Пробую по расширению...", file=sys.stderr)
        _, extension = os.path.splitext(file_path)
        
        # Fallback по расширению
        if extension.lower() == '.pdf':
            return extract_text_from_pdf(file_path)
        elif extension.lower() == '.docx':
            return extract_text_from_docx(file_path)
        elif extension.lower() == '.doc':
            return process_complex_doc(file_path, tempfile.gettempdir())
        elif extension.lower() == '.rtf':
            return extract_text_from_rtf(file_path)
        elif extension.lower() in ['.xlsx', '.xls']:
            return extract_text_from_excel(file_path)
        elif extension.lower() == '.csv':
            return extract_text_from_csv(file_path)
        elif extension.lower() == '.txt':
            return extract_text_from_txt(file_path)
        elif extension.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return extract_text_from_image(file_path)
        else:
            return f"ERROR: Неподдерживаемый или неизвестный тип файла: {file_type}"

def main():
    """Main function - processes a single file and returns JSON result"""
    if len(sys.argv) != 3:
        result = {
            "error": "Usage: python simple_text_extractor.py <file_path> <mime_type>",
            "success": False
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    file_path = sys.argv[1]
    mime_type = sys.argv[2]
    
    if not os.path.exists(file_path):
        result = {
            "error": f"File not found: {file_path}",
            "success": False
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    try:
        # Detect file type if not provided
        if not mime_type or mime_type == 'unknown':
            mime_type = detect_file_type(file_path)
        
        # Extract text based on file type
        if mime_type.startswith('application/pdf'):
            text = extract_text_from_pdf(file_path)
        elif mime_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            text = extract_text_from_docx(file_path)
        elif mime_type.startswith('application/msword'):
            print(f"INFO: Routing .doc file to complex processor (LibreOffice first)...", file=sys.stderr)
            text = process_complex_doc(file_path, tempfile.gettempdir())
        elif mime_type.startswith('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet') or \
             mime_type.startswith('application/vnd.ms-excel'):
            text = extract_text_from_excel(file_path)
        elif mime_type.startswith('text/csv'):
            text = extract_text_from_csv(file_path)  # Use pandas for CSV
        elif mime_type.startswith('text/'):
            text = extract_text_from_txt(file_path)
        elif mime_type.startswith('image/'):
            text = extract_text_from_image(file_path)  # pytesseract for images
        elif mime_type.startswith('application/rtf'):
            text = extract_text_from_rtf(file_path)  # RTF files
        else:
            # Fallback based on file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.pdf':
                text = extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                text = extract_text_from_docx(file_path)
            elif file_ext == '.doc':
                print(f"INFO: Routing .doc file to complex processor (LibreOffice first)...", file=sys.stderr)
                text = process_complex_doc(file_path, tempfile.gettempdir())
            elif file_ext in ['.xlsx', '.xls']:
                text = extract_text_from_excel(file_path)
            elif file_ext == '.csv':
                text = extract_text_from_csv(file_path)  # Use pandas for CSV
            elif file_ext == '.txt':
                text = extract_text_from_txt(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
                text = extract_text_from_image(file_path)  # EasyOCR for images
            elif file_ext == '.rtf':
                text = extract_text_from_rtf(file_path)  # RTF files
            else:
                result = {
                    "error": f"Unsupported file type: {mime_type}",
                    "success": False
                }
                print(json.dumps(result, ensure_ascii=False))
                sys.exit(1)
        
        # Return result
        result = {
            "success": True,
            "text": text,
            "text_length": len(text),
            "file_type": mime_type,
            "filename": os.path.basename(file_path)
        }
        
        # Use sys.stdout.buffer.write to handle Unicode properly
        json_output = json.dumps(result, ensure_ascii=False, default=str)
        sys.stdout.buffer.write(json_output.encode('utf-8'))
        sys.stdout.buffer.flush()
    
    except Exception as e:
        result = {
            "error": f"Error processing file: {e}",
            "success": False
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)

if __name__ == "__main__":
    main()
