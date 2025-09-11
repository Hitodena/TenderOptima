#!/usr/bin/env python3
"""
ENHANCED Russian Text Extraction for SupplierFinder Analysis
Specialized optimization for Russian document processing with improved OCR and encoding
"""

import sys
import os
import re
import tempfile
import chardet
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    from docx import Document
    import pandas as pd
    import openpyxl
except ImportError as e:
    print(f"Missing required package: {e}", file=sys.stderr)
    sys.exit(1)

def detect_text_encoding(file_path):
    """Enhanced encoding detection for Russian text"""
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
        
    # Prioritize Russian encodings
    detected_encoding = result.get('encoding', 'utf-8')
    if detected_encoding:
        detected_encoding = detected_encoding.lower()
    else:
        detected_encoding = 'utf-8'
    confidence = result.get('confidence', 0)
    
    # If confidence is low, try Russian-specific encodings
    if confidence < 0.7:
        russian_encodings = ['cp1251', 'koi8-r', 'iso-8859-5', 'utf-8']
        for encoding in russian_encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as test_file:
                    sample = test_file.read(1000)
                    # Check for Russian characters
                    if re.search(r'[а-яё]', sample, re.IGNORECASE):
                        return encoding
            except:
                continue
    
    return detected_encoding

def detect_sections_russian_enhanced(text):
    """Enhanced section detection optimized for Russian documents"""
    sections = {}
    current_section = "general"
    
    # RUSSIAN-OPTIMIZED section patterns with Cyrillic support
    section_patterns = [
        # PRIMARY: Multi-level numbered sections with Russian text
        (r'^\s*(\d+)\.(\d+)\.(\d+)[\s\.]*([А-ЯЁ][а-яё\s,\-:()]{3,})', 'numbered_subsection_L3'),
        (r'^\s*(\d+)\.(\d+)[\s\.]*([А-ЯЁ][а-яё\s,\-:()]{3,})', 'numbered_subsection_L2'),
        (r'^\s*(\d+)[\s\.]*([А-ЯЁ][а-яё\s,\-:()]{3,})', 'numbered_section_L1'),
        
        # SECONDARY: Section numbers with optional dots and spaces
        (r'^\s*(\d+)\.(\d+)\.(\d+)[\s\.]*$', 'section_number_L3'),
        (r'^\s*(\d+)\.(\d+)[\s\.]*$', 'section_number_L2'),
        (r'^\s*(\d+)[\s\.]*$', 'section_number_L1'),
        
        # TERTIARY: Russian section titles
        (r'^\s*(Раздел\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'section_title'),
        (r'^\s*(Глава\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'chapter_title'),
        (r'^\s*(Пункт\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'point_title'),
        (r'^\s*(Подраздел\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'subsection_title'),
        
        # QUATERNARY: Russian technical content patterns
        (r'^\s*(Техническ[а-я]*\s+характеристик[а-я]*)', 'technical_specifications'),
        (r'^\s*(Комплект[а-я]*\s+оборудован[а-я]*)', 'equipment'),
        (r'^\s*(Комплект[а-я]*\s+технологическ[а-я]*)', 'equipment'),
        (r'^\s*(Систем[а-я]*\s+автоматизац[а-я]*)', 'automation_control'),
        (r'^\s*(Систем[а-я]*\s+контрол[а-я]*)', 'automation_control'),
        (r'^\s*(Общ[а-я]*\s+требован[а-я]*)', 'general_requirements'),
        (r'^\s*(Безопасност[а-я]*)', 'safety'),
        (r'^\s*(Автоматическ[а-я]*\s+[а-я]*)', 'automation_equipment'),
        (r'^\s*(Электр[а-я]*\s+[а-я]*)', 'electrical'),
        (r'^\s*(Механическ[а-я]*\s+[а-я]*)', 'mechanical'),
        
        # Equipment-specific Russian patterns
        (r'^\s*([А-ЯЁ][а-яё]*[-]*[а-яё]*\s+плавитель)', 'melter_equipment'),
        (r'^\s*([А-ЯЁ][а-яё]*[-]*[а-яё]*\s+экструдер)', 'extruder_equipment'),
        (r'^\s*([А-ЯЁ][а-яё]*[-]*[а-яё]*\s+формовщик)', 'former_equipment'),
        (r'^\s*([А-ЯЁ][а-яё]*[-]*[а-яё]*\s+установк[а-я]*)', 'installation_equipment'),
    ]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Enhanced pattern matching for Russian text
        section_detected = False
        section_title = ""
        section_type = "content"
        
        for pattern, pattern_type in section_patterns:
            match = re.match(pattern, line_stripped, re.IGNORECASE | re.MULTILINE)
            if match:
                section_title = match.group(0).strip()
                section_type = pattern_type
                
                # Enhanced section categorization with hierarchical numbering
                if 'numbered' in pattern_type:
                    groups = [g for g in match.groups() if g and g.isdigit()]
                    if groups:
                        current_section = f"section_{'_'.join(groups)}"
                    else:
                        current_section = f"section_{len(sections) + 1}"
                        
                elif 'section_number' in section_type:
                    groups = [g for g in match.groups() if g and g.isdigit()]
                    if groups:
                        current_section = f"section_{'_'.join(groups)}"
                    else:
                        current_section = f"section_{len(sections) + 1}"
                        
                elif any(word in line_stripped.lower() for word in ['техническ', 'характеристик']):
                    current_section = "technical_specifications"
                elif any(word in line_stripped.lower() for word in ['комплект', 'оборудован', 'технологическ']):
                    current_section = "equipment"
                elif any(word in line_stripped.lower() for word in ['автоматизац', 'контрол']):
                    current_section = "automation_control"
                elif any(word in line_stripped.lower() for word in ['безопасност']):
                    current_section = "safety"
                elif any(word in line_stripped.lower() for word in ['общ', 'требован']):
                    current_section = "general_requirements"
                else:
                    current_section = f"section_{len(sections) + 1}"
                
                section_detected = True
                break
        
        # Initialize section if new
        if current_section not in sections:
            # Extract numeric order for proper sorting
            numeric_order = extract_numeric_order(current_section)
            sections[current_section] = {
                'title': section_title if section_detected and section_title else f"Раздел {current_section}",
                'content': [],
                'section_type': section_type if section_detected and section_type else 'content',
                'numeric_order': numeric_order
            }
        
        sections[current_section]['content'].append(line)
    
    return sections

def extract_numeric_order(section_name):
    """Extract numeric ordering from section name for proper sorting"""
    if section_name.startswith('section_'):
        numbers = section_name.replace('section_', '').split('_')
        try:
            return [int(n) for n in numbers if n.isdigit()]
        except ValueError:
            return [999]
    return [0]

def sort_sections_properly(sections):
    """Sort sections by proper numerical order"""
    def sort_key(item):
        section_name, section_data = item
        return section_data.get('numeric_order', [999])
    
    return sorted(sections.items(), key=sort_key)

def extract_text_from_pdf_russian_enhanced(file_path):
    """Enhanced PDF extraction optimized for Russian documents"""
    text = ""
    
    try:
        # Try PyPDF2 first with better encoding handling
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Clean up encoding issues common in Russian PDFs
                        page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8')
                        text += f"\n[PAGE:{page_num}]\n{page_text}\n"
                except Exception as e:
                    print(f"Error extracting page {page_num}: {e}", file=sys.stderr)
                    continue
        
        # Enhanced OCR fallback with Russian prioritization
        if not text.strip() or len(text) < 100:
            print("Minimal text found via PyPDF2, using enhanced OCR for Russian...", file=sys.stderr)
            try:
                images = convert_from_path(file_path, dpi=300)  # Higher DPI for better Russian text recognition
                for page_num, image in enumerate(images, 1):
                    # Enhanced OCR with Russian language prioritization
                    custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯабвгдеёжзийклмнопрстуфхцчшщъыьэюя0123456789.,()-/\\ '
                    ocr_text = pytesseract.image_to_string(
                        image, 
                        lang='rus',  # Prioritize Russian only
                        config=custom_config
                    )
                    if ocr_text.strip():
                        text += f"\n[PAGE:{page_num}]\n{ocr_text}\n"
            except Exception as e:
                print(f"OCR extraction failed: {e}", file=sys.stderr)
    
    except Exception as e:
        print(f"Error extracting from PDF: {e}", file=sys.stderr)
        return ""
    
    # Enhanced section detection for Russian text
    if text.strip():
        sections = detect_sections_russian_enhanced(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_docx_russian_enhanced(file_path):
    """Enhanced DOCX extraction optimized for Russian documents"""
    text = ""
    page_num = 1
    
    try:
        doc = Document(file_path)
        
        paragraph_texts = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Clean up encoding issues
                para_text = para_text.encode('utf-8', errors='ignore').decode('utf-8')
                paragraph_texts.append(para_text)
        
        for i, para_text in enumerate(paragraph_texts):
            if i > 0 and i % 50 == 0:
                page_num += 1
                text += f"\n[PAGE:{page_num}]\n"
            text += para_text + "\n"
        
        # Process tables with Russian text handling
        for table in doc.tables:
            text += f"\n[TABLE]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cell_text = cell_text.encode('utf-8', errors='ignore').decode('utf-8')
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += f"[/TABLE]\n"
        
    except Exception as e:
        print(f"Error extracting from DOCX: {e}", file=sys.stderr)
        return ""
    
    # Enhanced section detection for Russian text
    if text.strip():
        sections = detect_sections_russian_enhanced(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_txt_russian_enhanced(file_path):
    """Enhanced TXT extraction optimized for Russian documents"""
    text = ""
    
    try:
        # Enhanced encoding detection for Russian text
        encoding = detect_text_encoding(file_path)
        print(f"Detected encoding: {encoding}", file=sys.stderr)
        
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
        except UnicodeDecodeError:
            # Fallback encodings for Russian text
            russian_encodings = ['cp1251', 'koi8-r', 'iso-8859-5', 'utf-8', 'cp866']
            for fallback_encoding in russian_encodings:
                try:
                    with open(file_path, 'r', encoding=fallback_encoding) as file:
                        text = file.read()
                        print(f"Successfully read with fallback encoding: {fallback_encoding}", file=sys.stderr)
                        break
                except UnicodeDecodeError:
                    continue
    except Exception as e:
        print(f"Error reading text file: {e}", file=sys.stderr)
        return ""
    
    # Enhanced section detection for Russian text
    if text.strip():
        sections = detect_sections_russian_enhanced(text)
        sorted_sections = sort_sections_properly(sections)
        sectioned_text = ""
        
        for section_name, section_data in sorted_sections:
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += f"[NUMERIC_ORDER:{'.'.join(map(str, section_data['numeric_order']))}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_text_russian_enhanced.py <file_path> <mime_type>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    mime_type = sys.argv[2]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    # Route to appropriate extraction method
    if mime_type == 'application/pdf':
        extracted_text = extract_text_from_pdf_russian_enhanced(file_path)
    elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        extracted_text = extract_text_from_docx_russian_enhanced(file_path)
    elif mime_type == 'text/plain':
        extracted_text = extract_text_from_txt_russian_enhanced(file_path)
    else:
        # Fallback to optimized extractor for other formats
        from extract_text_optimized import extract_text_from_pdf_optimized, extract_text_from_docx_optimized, extract_text_from_txt_optimized
        if 'pdf' in mime_type:
            extracted_text = extract_text_from_pdf_optimized(file_path)
        elif 'word' in mime_type or 'document' in mime_type:
            extracted_text = extract_text_from_docx_optimized(file_path)
        else:
            extracted_text = extract_text_from_txt_optimized(file_path)
    
    print(extracted_text)

if __name__ == "__main__":
    main()