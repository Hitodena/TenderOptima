#!/usr/bin/env python3
"""
Section-Aware Text Extraction for SupplierFinder Analysis
Enhanced text extraction with section detection and structure preservation
"""

import sys
import os
import re
import tempfile
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    from docx import Document
    import pandas as pd
    import openpyxl
except ImportError as e:
    print(f"Missing required package: {e}", file=sys.stderr)
    sys.exit(1)

def detect_sections(text):
    """Detect document sections and structure for section-aware analysis"""
    sections = {}
    current_section = "general"
    
    # Enhanced section patterns for technical documents - IMPROVED ORDER AND DETECTION
    section_patterns = [
        # PRIMARY: Precise numbered sections (1., 2., 3., etc.)
        (r'^\s*(\d+)\.(\d+)?\.?(\d+)?\s+([А-ЯЁA-Z][а-яё\s,\-a-zA-Z]{4,})', 'numbered_subsection'),
        (r'^\s*(\d+)\.\s*([А-ЯЁA-Z][а-яё\s,\-a-zA-Z]{5,})', 'numbered_section'),
        (r'^\s*(\d+)\.?\s*$', 'section_number_only'),
        
        # SECONDARY: Formal section titles
        (r'^\s*(Раздел\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'section_title'),
        (r'^\s*(Глава\s+\d+[.\s]*[А-ЯЁ][а-яё\s]*)', 'chapter_title'),
        
        # Technical content patterns
        (r'^\s*(Техническ[а-я]*\s+характеристик[а-я]*)', 'technical_specifications'),
        (r'^\s*(Комплект[а-я]*\s+оборудован[а-я]*)', 'equipment'),
        (r'^\s*(Комплект[а-я]*\s+технологическ[а-я]*)', 'equipment'),
        (r'^\s*(Систем[а-я]*\s+автоматизац[а-я]*)', 'automation_control'),
        (r'^\s*(Систем[а-я]*\s+контрол[а-я]*)', 'automation_control'),
        (r'^\s*(Общ[а-я]*\s+требован[а-я]*)', 'general_requirements'),
        (r'^\s*(Безопасност[а-я]*)', 'safety'),
        (r'^\s*(Контрол[а-я]*)', 'control'),
        (r'^\s*(Электронн[а-я]*\s+модул[а-я]*)', 'electronic_modules'),
        (r'^\s*(Модул[а-я]*\s+управлен[а-я]*)', 'control_modules'),
        (r'^\s*(Блок[а-я]*\s+управлен[а-я]*)', 'control_units'),
        (r'^\s*(Датчик[а-я]*)', 'sensors'),
        (r'^\s*(Исполнительн[а-я]*\s+механизм[а-я]*)', 'actuators'),
        
        # Equipment specific patterns
        (r'^\s*(Насос[а-я]*)', 'pumps'),
        (r'^\s*(Двигател[а-я]*)', 'motors'),
        (r'^\s*(Редуктор[а-я]*)', 'gearboxes'),
        (r'^\s*(Экструдер[а-я]*)', 'extruders'),
        (r'^\s*(Формовщик[а-я]*)', 'molding_equipment'),
        (r'^\s*(Установк[а-я]*)', 'installations'),
        
        # English patterns
        (r'^\s*(\d+)\.?\s*([A-Z][a-z\s,\-]{5,})', 'numbered_section_en'),
        (r'^\s*(Technical\s+Specifications?)', 'technical_specifications'),
        (r'^\s*(Equipment\s+Set)', 'equipment'),
        (r'^\s*(Automation\s+System)', 'automation_control'),
        (r'^\s*(Control\s+System)', 'automation_control'),
        (r'^\s*(General\s+Requirements)', 'general_requirements'),
        (r'^\s*(Safety\s+Requirements)', 'safety'),
        (r'^\s*(Electronic\s+Modules?)', 'electronic_modules'),
        (r'^\s*(Control\s+Modules?)', 'control_modules'),
    ]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Check for section headers
        section_detected = False
        for pattern, section_type in section_patterns:
            match = re.match(pattern, line_stripped, re.IGNORECASE | re.MULTILINE)
            if match:
                section_key = match.group(1).lower() if len(match.groups()) >= 1 else section_type
                section_title = match.group(0).strip()
                
                # Intelligent section categorization
                if any(word in line_stripped.lower() for word in ['техническ', 'характеристик', 'technical', 'specification']):
                    current_section = "technical_specifications"
                elif any(word in line_stripped.lower() for word in ['комплект', 'оборудован', 'equipment', 'технологическ']):
                    current_section = "equipment"
                elif any(word in line_stripped.lower() for word in ['автоматизац', 'automation', 'контрол', 'control']):
                    current_section = "automation_control"
                elif any(word in line_stripped.lower() for word in ['безопасност', 'safety']):
                    current_section = "safety"
                elif any(word in line_stripped.lower() for word in ['общ', 'требован', 'general', 'requirement']):
                    current_section = "general_requirements"
                elif any(word in line_stripped.lower() for word in ['электронн', 'модул', 'electronic', 'module']):
                    current_section = "electronic_modules"
                elif any(word in line_stripped.lower() for word in ['насос', 'pump']):
                    current_section = "pumps"
                elif any(word in line_stripped.lower() for word in ['двигател', 'motor']):
                    current_section = "motors"
                elif any(word in line_stripped.lower() for word in ['экструдер', 'extruder']):
                    current_section = "extruders"
                elif any(word in line_stripped.lower() for word in ['формовщик', 'molding']):
                    current_section = "molding_equipment"
                else:
                    # Extract section number for numbered sections - IMPROVED NUMERICAL DETECTION
                    number_match = re.match(r'^\s*(\d+)\.?(\d+)?\.?(\d+)?', line_stripped)
                    if number_match:
                        # Build hierarchical section number (e.g., "4.1.2")
                        section_parts = [number_match.group(1)]
                        if number_match.group(2):
                            section_parts.append(number_match.group(2))
                        if number_match.group(3):
                            section_parts.append(number_match.group(3))
                        current_section = f"section_{'_'.join(section_parts)}"
                    else:
                        current_section = f"section_{len(sections) + 1}"
                
                if current_section not in sections:
                    sections[current_section] = {
                        'title': section_title,
                        'content': [],
                        'start_line': i,
                        'section_type': section_type
                    }
                
                section_detected = True
                break
        
        # Add content to current section
        if current_section not in sections:
            sections[current_section] = {
                'title': 'General Content',
                'content': [],
                'start_line': 0,
                'section_type': 'general'
            }
        
        sections[current_section]['content'].append(line)
    
    return sections

def extract_text_from_pdf(file_path):
    """Extract text from PDF with section detection"""
    text = ""
    
    try:
        # First try PyPDF2 for text-based PDFs
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n[PAGE:{page_num}]\n{page_text}\n"
        
        # If no text extracted, use OCR
        if not text.strip():
            images = convert_from_path(file_path)
            for page_num, image in enumerate(images, 1):
                ocr_text = pytesseract.image_to_string(image, lang='rus+eng')
                if ocr_text.strip():
                    text += f"\n[PAGE:{page_num}]\n{ocr_text}\n"
    
    except Exception as e:
        print(f"Error extracting from PDF: {e}", file=sys.stderr)
    
    # Add section detection
    if text.strip():
        sections = detect_sections(text)
        sectioned_text = ""
        
        for section_name, section_data in sections.items():
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX with section detection"""
    text = ""
    page_num = 1
    
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                if len(text.split('\n')) % 50 == 0 and len(text) > 0:
                    page_num += 1
                    text += f"\n[PAGE:{page_num}]\n"
                text += para_text + "\n"
        
        # Extract tables
        for table in doc.tables:
            text += f"\n[TABLE]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += f"[/TABLE]\n"
        
        # Extract headers and footers
        try:
            for section in doc.sections:
                header = section.header
                footer = section.footer
                
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            text += f"[HEADER] {para.text.strip()}\n"
                
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text += f"[FOOTER] {para.text.strip()}\n"
        except:
            pass
        
    except Exception as e:
        print(f"Error extracting from DOCX: {e}", file=sys.stderr)
        return ""
    
    # Add section detection
    if text.strip():
        sections = detect_sections(text)
        sectioned_text = ""
        
        for section_name, section_data in sections.items():
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def extract_text_from_txt(file_path):
    """Extract text from plain text files with section detection"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='cp1251') as file:
                text = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='iso-8859-1') as file:
                text = file.read()
    
    # Add section detection
    if text.strip():
        sections = detect_sections(text)
        sectioned_text = ""
        
        for section_name, section_data in sections.items():
            sectioned_text += f"\n[SECTION:{section_name.upper()}]\n"
            sectioned_text += f"[SECTION_TITLE:{section_data['title']}]\n"
            sectioned_text += f"[SECTION_TYPE:{section_data['section_type']}]\n"
            sectioned_text += '\n'.join(section_data['content'])
            sectioned_text += f"\n[/SECTION:{section_name.upper()}]\n"
        
        return sectioned_text if sectioned_text.strip() else text
    
    return text

def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_text_with_sections.py <file_path> <mime_type>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    mime_type = sys.argv[2]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}", file=sys.stderr)
        sys.exit(1)
    
    try:
        if mime_type.startswith('application/pdf'):
            text = extract_text_from_pdf(file_path)
        elif mime_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
            text = extract_text_from_docx(file_path)
        elif mime_type.startswith('text/'):
            text = extract_text_from_txt(file_path)
        else:
            # Fallback based on file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.pdf':
                text = extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                text = extract_text_from_txt(file_path)
            else:
                print(f"Unsupported file type: {mime_type}", file=sys.stderr)
                sys.exit(1)
        
        print(text)
    
    except Exception as e:
        print(f"Error processing file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()